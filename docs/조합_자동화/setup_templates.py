"""
setup_templates.py
────────────────────────────────────────────────────────
[1단계: 최초 1회 실행]
3호 조합 완성본 폴더에서 templates/ 폴더로 파일을 복사하면서
알려진 변수값을 {{마커}} 형태로 자동 치환합니다.

실행 방법:
  python setup_templates.py --source "C:/경로/[트리거메디테크3호조합]_비식별"
"""

import os
import sys
import json
import shutil
import argparse
import win32com.client
from docx import Document
from pathlib import Path


# ─────────────────────────────────────────────────
# 3호 조합 알려진 값 → 마커 매핑
# (더 긴 문자열을 먼저 치환해야 부분치환 오류 없음)
# ─────────────────────────────────────────────────

KNOWN_VALUES = [
    # (실제값, 마커키)  — 길이 긴 순으로 정렬 중요
    ("트리거 메디테크 3호 조합",        "조합명"),
    ("트리거메디테크3호조합",            "조합명_파일명"),
    ("트리거투자파트너스 유한회사",      "업무집행조합원_정식"),
    ("트리거투자파트너스(유)",           "업무집행조합원_약칭"),

    ("금이십이억칠천오백만원",           "총출자금_한글"),
    ("₩2,275,000,000",                   "총출자금_기호"),
    ("2,275,000,000",                    "총출자금_숫자"),
    ("2,275,000,000 원",                 "총출자금_숫자"),   # 공백 포함 버전도
    ("2,275,000,000원",                  "총출자금_숫자"),
    ("2,275좌",                          "총출자좌수"),
    ("2,275",                            "총출자좌수"),

    ("2025년 10월 24일(금요일) 오전 10시", "결성총회일시_풀"),
    ("2025년 10월 24일(금요일)",           "결성총회일시_풀"),  # 축약 버전
    ("2025년 10월 24일",                   "결성총회일_날짜"),
    ("2025. 10. 20",                       "소집통지날짜_약식"),
    ("2025년 10월 15일",                   "소집통지날짜"),
    ("2025년   10월  27일",                "등록신청일"),

    ("2025년 10월 24일(금) 오전 10시까지", "납입기한"),

    ("서울특별시 마포구 양화로7길 70, 5층(서교동)", "사업장주소_정식"),
    ("서울 마포구 양화로7길 70,컬처큐브 5층",      "사업장주소_약식"),

    ("국민은행) 003190-15-050045",   "납입계좌번호"),  # 괄호 포함 버전
    ("003190-15-050045",             "납입계좌번호"),

    ("태성회계법인",   "외부감사인"),
    ("유안타증권",     "수탁기관"),

    ("786-88-02871",  "사업자등록번호"),
    ("164714-0005810","법인등록번호"),
    ("479-80-03340",  "고유번호"),

    ("트리거-2025-23호", "문서번호_통지"),
    ("트리거-2025-28호", "문서번호_등록"),

    ("hwh@triggerip.com",  "담당자이메일"),
    ("010-9473-3142",      "담당자연락처"),
    ("홍운학",             "담당자명"),
    ("서원일",             "대표이사"),

    ("2025.09.19 ~ 2027.01.19",  "임대차_기간"),
    ("2027.01.19",               "임대차_종료"),
    ("2025.09.19",               "임대차_시작"),
    ("111.83㎡",                 "임대차_면적"),

    ("02-2038-2456",  "전화번호"),
    ("02-6953-2456",  "팩스번호"),

    ("19명",  "조합원총수"),
    ("18명",  "유한책임조합원수"),
    ("5년",   "존속기간"),
]

# 길이 내림차순 정렬 (긴 문자열 먼저 치환)
KNOWN_VALUES.sort(key=lambda x: len(x[0]), reverse=True)


def build_replacement_map():
    """실제값 → {{마커}} 딕셔너리 반환"""
    return {val: f"{{{{{key}}}}}" for val, key in KNOWN_VALUES}


# ─────────────────────────────────────────────────
# HWP 마커 삽입
# ─────────────────────────────────────────────────

class HwpMarker:
    def __init__(self):
        self.hwp = None

    def __enter__(self):
        self.hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
        try:
            self.hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        except Exception:
            pass
        try:
            self.hwp.XHwpWindows.Item(0).Visible = False
        except Exception:
            pass
        return self

    def __exit__(self, *args):
        if self.hwp:
            try:
                self.hwp.Quit()
            except Exception:
                pass

    def mark_file(self, src: str, dst: str, replacements: dict) -> tuple[bool, int]:
        """src → dst 로 치환하여 저장. (성공여부, 치환횟수) 반환"""
        abs_src = str(Path(src).resolve())
        abs_dst = str(Path(dst).resolve())
        Path(abs_dst).parent.mkdir(parents=True, exist_ok=True)

        if not self.hwp.Open(abs_src, "HWP", "forceopen:true"):
            return False, 0

        total_replaced = 0
        for find_str, replace_str in replacements.items():
            count = self._find_replace(find_str, replace_str)
            total_replaced += count

        self.hwp.SaveAs(abs_dst, "HWP")
        self.hwp.Clear(1)
        return True, total_replaced

    def _find_replace(self, find_str, replace_str) -> int:
        try:
            act = self.hwp.CreateAction("AllReplace")
            pset = act.CreateSet()
            act.GetDefault(pset)
            pset.SetItem("FindString", find_str)
            pset.SetItem("ReplaceString", replace_str)
            pset.SetItem("IgnoreCase", 0)
            pset.SetItem("WholeWordOnly", 0)
            pset.SetItem("ReplaceMode", 1)
            result = act.Execute(pset)
            return 1 if result else 0
        except Exception:
            return 0


# ─────────────────────────────────────────────────
# DOCX 마커 삽입
# ─────────────────────────────────────────────────

def mark_docx(src: str, dst: str, replacements: dict) -> tuple[bool, int]:
    try:
        Path(dst).parent.mkdir(parents=True, exist_ok=True)
        doc = Document(src)
        total = 0

        def _replace_para(para):
            nonlocal total
            full = "".join(r.text for r in para.runs)
            new_full = full
            for find_str, replace_str in replacements.items():
                if find_str in new_full:
                    new_full = new_full.replace(find_str, replace_str)
                    total += 1
            if new_full != full and para.runs:
                para.runs[0].text = new_full
                for r in para.runs[1:]:
                    r.text = ""

        for para in doc.paragraphs:
            _replace_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_para(para)

        doc.save(dst)
        return True, total
    except Exception as e:
        print(f"  DOCX 오류: {e}")
        return False, 0


# ─────────────────────────────────────────────────
# 메인 실행
# ─────────────────────────────────────────────────

def setup_templates(source_dir: str, template_dir: str, log_callback=None):
    """
    source_dir: 3호 조합 완성본 폴더
    template_dir: 결과 templates/ 폴더
    """
    src = Path(source_dir)
    tpl = Path(template_dir)
    tpl.mkdir(parents=True, exist_ok=True)

    replacements = build_replacement_map()

    hwp_files = list(src.rglob("*.hwp"))
    docx_files = list(src.rglob("*.docx"))
    other_files = [f for f in src.rglob("*")
                   if f.is_file() and f.suffix.lower() not in (".hwp", ".docx")]

    total = len(hwp_files) + len(docx_files) + len(other_files)
    current = 0
    summary = {"hwp": 0, "docx": 0, "copy": 0, "failed": []}

    # HWP 처리
    with HwpMarker() as marker:
        for f in hwp_files:
            current += 1
            rel = f.relative_to(src)
            dst = tpl / rel
            if log_callback:
                log_callback(current, total, f"HWP 마킹: {f.name}")
            ok, count = marker.mark_file(str(f), str(dst), replacements)
            if ok:
                summary["hwp"] += 1
                if log_callback:
                    log_callback(current, total, f"  → {count}개 치환 완료")
            else:
                summary["failed"].append(str(rel))

    # DOCX 처리
    for f in docx_files:
        current += 1
        rel = f.relative_to(src)
        dst = tpl / rel
        if log_callback:
            log_callback(current, total, f"DOCX 마킹: {f.name}")
        ok, count = mark_docx(str(f), str(dst), replacements)
        if ok:
            summary["docx"] += 1
        else:
            summary["failed"].append(str(rel))

    # 나머지 파일 복사
    for f in other_files:
        current += 1
        rel = f.relative_to(src)
        dst = tpl / rel
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(f, dst)
        summary["copy"] += 1
        if log_callback:
            log_callback(current, total, f"복사: {f.name}")

    return summary


# ─────────────────────────────────────────────────
# CLI 진입점
# ─────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="3호 조합 → 템플릿 마킹 도구")
    parser.add_argument(
        "--source",
        default=r"C:\Users\1llal\Desktop\[트리거메디테크3호조합]_비식별",
        help="3호 조합 완성본 폴더 경로"
    )
    parser.add_argument(
        "--output",
        default=str(Path(__file__).parent / "templates"),
        help="templates 폴더 경로"
    )
    args = parser.parse_args()

    print(f"소스 폴더: {args.source}")
    print(f"출력 폴더: {args.output}")
    print()

    def on_log(cur, tot, msg):
        print(f"[{cur:>3}/{tot}] {msg}")

    result = setup_templates(args.source, args.output, log_callback=on_log)

    print()
    print("=" * 50)
    print(f"HWP 마킹 완료: {result['hwp']}개")
    print(f"DOCX 마킹 완료: {result['docx']}개")
    print(f"그대로 복사: {result['copy']}개")
    if result["failed"]:
        print(f"실패: {len(result['failed'])}개")
        for f in result["failed"]:
            print(f"  - {f}")
    print()
    print("템플릿 준비 완료! 이제 app.py 를 실행하세요.")
