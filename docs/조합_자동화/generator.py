"""
generator.py
HWP COM + python-docx 기반 조합 서류 자동 생성 엔진
"""

import os
import json
import shutil
import win32com.client
from docx import Document
from pathlib import Path


# ─────────────────────────────────────────────────
# HWP COM 래퍼
# ─────────────────────────────────────────────────

class HwpAutomation:
    def __init__(self, visible=False):
        self.hwp = None
        self.visible = visible

    def __enter__(self):
        self.hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
        try:
            # 보안 모듈 등록 (파일 경로 접근 허용)
            self.hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        except Exception:
            pass
        if not self.visible:
            try:
                self.hwp.XHwpWindows.Item(0).Visible = False
            except Exception:
                pass
        return self

    def __exit__(self, *args):
        if self.hwp:
            try:
                self.hwp.Quit()
            except Exception:
                pass
            self.hwp = None

    def replace_in_file(self, src_path: str, dst_path: str, replacements: dict) -> tuple[bool, list]:
        """
        src_path 의 HWP 파일에서 replacements 의 키를 값으로 치환하여 dst_path 에 저장.
        반환: (성공여부, 치환 미발견 키 목록)
        """
        abs_src = str(Path(src_path).resolve())
        abs_dst = str(Path(dst_path).resolve())

        # dst 폴더 생성
        Path(abs_dst).parent.mkdir(parents=True, exist_ok=True)

        # 파일 열기
        if not self.hwp.Open(abs_src, "HWP", "forceopen:true"):
            return False, list(replacements.keys())

        not_found = []
        for marker, value in replacements.items():
            marker_str = f"{{{{{marker}}}}}"   # {{키}}
            found = self._find_replace(marker_str, str(value))
            if not found:
                not_found.append(marker)

        # 저장 (HWP 형식 유지)
        self.hwp.SaveAs(abs_dst, "HWP")
        self.hwp.Clear(1)  # 변경사항 폐기 후 닫기
        return True, not_found

    def _find_replace(self, find_str: str, replace_str: str) -> bool:
        """HWP AllReplace 액션으로 치환. 발견된 경우 True."""
        try:
            act = self.hwp.CreateAction("AllReplace")
            pset = act.CreateSet()
            act.GetDefault(pset)
            pset.SetItem("FindString", find_str)
            pset.SetItem("ReplaceString", replace_str)
            pset.SetItem("IgnoreCase", 0)
            pset.SetItem("WholeWordOnly", 0)
            pset.SetItem("AllWordForms", 0)
            pset.SetItem("SeveralWords", 0)
            pset.SetItem("UseWildCards", 0)
            pset.SetItem("FindRegExp", 0)
            pset.SetItem("ReplaceMode", 1)    # 모두 바꾸기
            pset.SetItem("FindJaso", 0)
            pset.SetItem("HanjaFromHangul", 0)
            result = act.Execute(pset)
            return result != 0
        except Exception:
            return False


# ─────────────────────────────────────────────────
# DOCX 치환
# ─────────────────────────────────────────────────

def replace_in_docx(src_path: str, dst_path: str, replacements: dict) -> tuple[bool, list]:
    """DOCX 파일의 마커를 치환하여 저장."""
    try:
        Path(dst_path).parent.mkdir(parents=True, exist_ok=True)
        doc = Document(src_path)
        not_found = []
        found_keys = set()

        def _replace_in_runs(runs):
            # 런(run) 단위로 치환 — 마커가 한 런에 걸쳐있을 경우 처리
            full_text = "".join(r.text for r in runs)
            for marker, value in replacements.items():
                marker_str = f"{{{{{marker}}}}}"
                if marker_str in full_text:
                    full_text = full_text.replace(marker_str, str(value))
                    found_keys.add(marker)
            # 치환된 텍스트를 첫 런에 넣고 나머지 런은 비움
            if runs:
                runs[0].text = full_text
                for r in runs[1:]:
                    r.text = ""

        # 본문 단락
        for para in doc.paragraphs:
            _replace_in_runs(para.runs)

        # 표 안 셀
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_runs(para.runs)

        not_found = [k for k in replacements if k not in found_keys]
        doc.save(dst_path)
        return True, not_found
    except Exception as e:
        return False, [str(e)]


# ─────────────────────────────────────────────────
# 메인 생성 함수
# ─────────────────────────────────────────────────

def generate_documents(
    variables: dict,
    template_dir: str,
    output_dir: str,
    progress_callback=None
) -> dict:
    """
    template_dir 의 모든 HWP/DOCX 파일에 variables 를 치환하여
    output_dir/[조합명]/ 아래에 동일한 폴더 구조로 저장.

    progress_callback(current, total, message) 형태로 진행상황 전달.
    반환: { 'success': [...], 'failed': [...], 'warnings': [...] }
    """
    template_path = Path(template_dir)
    fund_name = variables.get("조합명", "새조합")
    output_base = Path(output_dir) / fund_name
    output_base.mkdir(parents=True, exist_ok=True)

    # 처리할 파일 목록 수집
    all_files = []
    for f in template_path.rglob("*"):
        if f.is_file() and f.suffix.lower() in (".hwp", ".docx"):
            all_files.append(f)
    # PDF/이미지 등은 그대로 복사
    copy_files = []
    for f in template_path.rglob("*"):
        if f.is_file() and f.suffix.lower() not in (".hwp", ".docx"):
            copy_files.append(f)

    total = len(all_files) + len(copy_files)
    results = {"success": [], "failed": [], "warnings": []}
    current = 0

    # ── HWP/DOCX 치환 처리 ──
    with HwpAutomation(visible=False) as hwp_auto:
        for f in all_files:
            current += 1
            rel = f.relative_to(template_path)
            dst = output_base / rel
            msg = f"처리 중: {f.name}"
            if progress_callback:
                progress_callback(current, total, msg)

            if f.suffix.lower() == ".hwp":
                ok, not_found = hwp_auto.replace_in_file(str(f), str(dst), variables)
            else:
                ok, not_found = replace_in_docx(str(f), str(dst), variables)

            if ok:
                results["success"].append(str(rel))
                if not_found:
                    results["warnings"].append(
                        f"{rel.name}: 미치환 마커 {not_found}"
                    )
            else:
                results["failed"].append(str(rel))

    # ── 나머지 파일 복사 ──
    for f in copy_files:
        current += 1
        rel = f.relative_to(template_path)
        dst = output_base / rel
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(f, dst)
        if progress_callback:
            progress_callback(current, total, f"복사: {f.name}")
        results["success"].append(str(rel) + " (복사)")

    if progress_callback:
        progress_callback(total, total, "완료!")
    return results


# ─────────────────────────────────────────────────
# 직접 실행 시 테스트
# ─────────────────────────────────────────────────

if __name__ == "__main__":
    base = Path(__file__).parent
    var_file = base / "variables" / "트리거메디테크3호_참고.json"

    with open(var_file, encoding="utf-8") as f:
        variables = json.load(f)
    # _comment 키 제거
    variables = {k: v for k, v in variables.items() if not k.startswith("_")}

    def on_progress(cur, tot, msg):
        print(f"[{cur}/{tot}] {msg}")

    results = generate_documents(
        variables=variables,
        template_dir=str(base / "templates"),
        output_dir=str(base / "output"),
        progress_callback=on_progress,
    )

    print("\n=== 결과 ===")
    print(f"성공: {len(results['success'])}개")
    print(f"실패: {len(results['failed'])}개")
    if results["warnings"]:
        print("경고:")
        for w in results["warnings"]:
            print(f"  - {w}")
    if results["failed"]:
        print("실패:")
        for ff in results["failed"]:
            print(f"  - {ff}")
