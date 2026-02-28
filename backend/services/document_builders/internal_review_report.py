from __future__ import annotations

import json
from collections import Counter
from datetime import date
from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from sqlalchemy.orm import Session

from models.fund import Fund
from models.internal_review import CompanyReview, InternalReview
from models.investment import Investment, PortfolioCompany

from .layout_utils import add_paragraph, create_base_document, set_cell_shading, set_cell_text, set_table_borders


def _fmt_date(value: date | None) -> str:
    if value is None:
        return "-"
    return value.isoformat()


def _fmt_money(value: float | int | None) -> str:
    if value is None:
        return "-"
    return f"{int(round(float(value))):,}원"


def _safe_text(value: str | None) -> str:
    return (value or "").strip() or "-"


def _load_attendees(value: str | None) -> list[str]:
    if not value:
        return []
    try:
        payload = json.loads(value)
    except json.JSONDecodeError:
        return []
    if not isinstance(payload, list):
        return []

    result: list[str] = []
    for row in payload:
        if isinstance(row, dict):
            name = row.get("name")
            role = row.get("role")
            if name and role:
                result.append(f"{name} ({role})")
            elif name:
                result.append(str(name))
        elif row:
            result.append(str(row))
    return result


def build_internal_review_report(internal_review_id: int, db: Session) -> bytes:
    """Build D-05 internal review integrated report docx bytes."""
    review = db.get(InternalReview, internal_review_id)
    if not review:
        raise LookupError("내부보고회 데이터를 찾을 수 없습니다.")

    fund = db.get(Fund, review.fund_id)
    if not fund:
        raise LookupError("조합 데이터를 찾을 수 없습니다.")

    company_reviews = (
        db.query(CompanyReview)
        .filter(CompanyReview.review_id == review.id)
        .order_by(CompanyReview.id.asc())
        .all()
    )

    rating_counter = Counter((row.asset_rating or "-") for row in company_reviews)
    impairment_count = sum(1 for row in company_reviews if (row.impairment_type or "none") != "none")

    doc = create_base_document()
    add_paragraph(
        doc,
        "내부보고회 통합보고서",
        size=17,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_paragraph(
        doc,
        f"{fund.name} · {review.year}년 {review.quarter}분기",
        size=11,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    summary_rows = [
        ("기준일", _fmt_date(review.reference_date)),
        ("보고회일", _fmt_date(review.review_date)),
        ("상태", _safe_text(review.status)),
        ("총 피투자사", f"{len(company_reviews)}개사"),
        ("등급 분포", ", ".join(f"{key}:{value}" for key, value in sorted(rating_counter.items())) or "-"),
        ("손상 검토 건수", f"{impairment_count}건"),
    ]
    summary_table = doc.add_table(rows=len(summary_rows), cols=2)
    summary_table.columns[0].width = Cm(5.0)
    summary_table.columns[1].width = Cm(11.0)
    set_table_borders(summary_table)
    for idx, (label, value) in enumerate(summary_rows):
        set_cell_text(summary_table.cell(idx, 0), label, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(summary_table.cell(idx, 1), value, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        if idx % 2 == 0:
            set_cell_shading(summary_table.cell(idx, 0), "F5F5F5")

    add_paragraph(doc, "", size=9, space_before=8)
    add_paragraph(doc, "피투자사별 상세", size=12, bold=True, space_after=4)

    for row in company_reviews:
        investment = db.get(Investment, row.investment_id)
        company = db.get(PortfolioCompany, investment.company_id) if investment else None
        company_name = company.name if company else f"Investment #{row.investment_id}"

        add_paragraph(doc, company_name, size=11, bold=True, space_before=6, space_after=2)
        info_table = doc.add_table(rows=8, cols=2)
        info_table.columns[0].width = Cm(5.0)
        info_table.columns[1].width = Cm(11.0)
        set_table_borders(info_table)

        info_rows = [
            ("분기 매출액", _fmt_money(row.quarterly_revenue)),
            ("분기 영업이익", _fmt_money(row.quarterly_operating_income)),
            ("분기 당기순이익", _fmt_money(row.quarterly_net_income)),
            ("총자산", _fmt_money(row.total_assets)),
            ("총부채", _fmt_money(row.total_liabilities)),
            ("자산건전성 등급", _safe_text(row.asset_rating)),
            ("손상유형", _safe_text(row.impairment_type)),
            ("투자의견", _safe_text(row.investment_opinion)),
        ]
        for idx, (label, value) in enumerate(info_rows):
            set_cell_text(info_table.cell(idx, 0), label, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(info_table.cell(idx, 1), value, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            if idx % 2 == 0:
                set_cell_shading(info_table.cell(idx, 0), "F9F9F9")

        add_paragraph(doc, f"주요 이슈: {_safe_text(row.key_issues)}", size=9, space_before=2)
        add_paragraph(doc, f"후속 조치: {_safe_text(row.follow_up_actions)}", size=9, space_before=1)
        add_paragraph(doc, f"의견: {_safe_text(row.investment_opinion)}", size=9, space_before=1)

    add_paragraph(doc, "", size=9, space_before=10)
    add_paragraph(doc, "준법감시인 의견", size=11, bold=True, space_after=2)
    add_paragraph(doc, _safe_text(review.compliance_opinion), size=10, space_after=6)
    add_paragraph(doc, f"준법감시인: {_safe_text(review.compliance_officer)}", size=10, space_after=6)

    attendees = _load_attendees(review.attendees_json)
    add_paragraph(doc, "참석자", size=11, bold=True, space_after=2)
    if attendees:
        for person in attendees:
            add_paragraph(doc, f"- {person}", size=10, space_before=1)
    else:
        add_paragraph(doc, "- 참석자 정보 없음", size=10)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
