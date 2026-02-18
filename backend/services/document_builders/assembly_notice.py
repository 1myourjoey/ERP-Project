from __future__ import annotations

from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH

from .layout_utils import add_paragraph, add_run, create_base_document, resolve_layout_scale

DEFAULT_AGENDAS = [
    "제1호 안건: 조합 규약 확인의 건",
    "제2호 안건: 자산보관/관리기관 운영방안 확인의 건",
    "제3호 안건: 회계감사인 선정의 건",
]


def _render_value(text: str, variables: dict) -> str:
    rendered = text
    for key, value in variables.items():
        if key.startswith("__"):
            continue
        rendered = rendered.replace("{{" + key + "}}", str(value))
    return rendered


def _as_dict(value) -> dict:
    return value if isinstance(value, dict) else {}


def _as_list(value) -> list:
    return value if isinstance(value, list) else []


def build_assembly_notice(variables: dict) -> BytesIO:
    """Build the formation assembly notice document."""
    custom = _as_dict(variables.get("__custom_data__"))
    layout_scale = resolve_layout_scale(custom)

    greeting = str(custom.get("greeting") or "조합원 여러분의 건승을 기원합니다.")
    regulation_article = str(custom.get("regulation_article") or "제15조")
    body_default = (
        "{{fund_name}} 규약 {{regulation_article}}에 따라 아래와 같이 결성총회를 소집하오니 "
        "서면결의로 의결권을 행사해 주시기 바랍니다."
    )
    body_text = str(custom.get("body_text") or body_default)
    agendas = _as_list(custom.get("agendas")) or DEFAULT_AGENDAS

    render_variables = dict(variables)
    render_variables["regulation_article"] = regulation_article

    doc = create_base_document(scale=layout_scale)

    def _p(*args, **kwargs):
        kwargs.setdefault("scale", layout_scale)
        return add_paragraph(*args, **kwargs)

    def _r(*args, **kwargs):
        kwargs.setdefault("scale", layout_scale)
        return add_run(*args, **kwargs)

    _p(doc, "[첨부 1] 결성총회 소집통지서", size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)
    _p(doc, str(variables.get("fund_name", "")), size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, "결성총회 소집통지서", size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    _p(doc, _render_value(greeting, render_variables), size=10, space_after=8)

    body_para = _p(doc, space_after=8)
    _r(body_para, _render_value(body_text, render_variables), size=10)

    _p(doc, "- 다 음 -", size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=8)

    _p(
        doc,
        f"1. 결성총회 일시 : {variables.get('assembly_date', '')} {variables.get('assembly_time', '오전 10시')}",
        size=10,
        space_before=2,
    )
    _p(doc, f"2. 결성총회 방법 : {variables.get('assembly_method', '서면결의')}", size=10, space_before=2)
    _p(doc, "3. 결의 목적사항", size=10, space_before=2, space_after=4)

    _p(doc, "안건", size=10, bold=True, space_before=4, space_after=4)
    for agenda in agendas:
        _p(doc, _render_value(str(agenda), render_variables), size=10, space_before=1)

    _p(doc, str(variables.get("document_date", "")), size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=18, space_after=8)
    _p(doc, str(variables.get("fund_name", "")), size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)
    _p(doc, f"업무집행조합원 {variables.get('gp_name', '')}", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=3)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer