from __future__ import annotations

from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from .layout_utils import (
    add_paragraph,
    create_base_document,
    resolve_layout_scale,
    set_cell_shading,
    set_cell_text,
    set_table_borders,
)

DEFAULT_INTRODUCTION = (
    "본인은 {{assembly_date}} 개최되는 {{fund_name}} 결성총회에 직접 참석하지 못하여, "
    "아래 안건에 대하여 서면으로 의결권을 행사합니다."
)

DEFAULT_AGENDAS = [
    "제1호 안건: 조합 규약 확인의 건",
    "제2호 안건: 자산보관/관리기관 운영방안 확인의 건",
    "제3호 안건: 회계감사인 선정의 건",
]

DEFAULT_VOTE_NOTE = "*찬성 또는 반대에 O/X로 표시해 주시기 바랍니다."


def _as_dict(value) -> dict:
    return value if isinstance(value, dict) else {}


def _as_list(value) -> list:
    return value if isinstance(value, list) else []


def _render_value(text: str, variables: dict) -> str:
    rendered = text
    for key, value in variables.items():
        if key.startswith("__"):
            continue
        rendered = rendered.replace("{{" + key + "}}", str(value))
    return rendered


def build_written_resolution(variables: dict) -> BytesIO:
    """Build a written resolution document."""
    custom = _as_dict(variables.get("__custom_data__"))
    layout_scale = resolve_layout_scale(custom)

    fund_name = str(variables.get("fund_name", "") or "")
    gp_name = str(variables.get("gp_name", "") or "")
    assembly_date = str(variables.get("assembly_date", "") or "")

    introduction_text = str(custom.get("introduction_text") or DEFAULT_INTRODUCTION)
    agendas = _as_list(custom.get("agendas")) or DEFAULT_AGENDAS
    vote_note = str(custom.get("vote_note") or DEFAULT_VOTE_NOTE)

    doc = create_base_document(scale=layout_scale)

    def _p(*args, **kwargs):
        kwargs.setdefault("scale", layout_scale)
        return add_paragraph(*args, **kwargs)

    def _c(*args, **kwargs):
        kwargs.setdefault("scale", layout_scale)
        return set_cell_text(*args, **kwargs)

    _p(doc, "[별첨] 서면결의서", size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    _p(doc, f"{fund_name} 업무집행조합원 귀중", size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
    if gp_name:
        _p(doc, f"({gp_name})", size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    _p(doc, _render_value(introduction_text, variables), size=10, space_after=10)

    vote_table = doc.add_table(rows=len(agendas) + 1, cols=3)
    set_table_borders(vote_table)
    vote_table.columns[0].width = Cm(10.0)
    vote_table.columns[1].width = Cm(2.5)
    vote_table.columns[2].width = Cm(2.5)

    headers = ["안건", "찬성", "반대"]
    for index, header in enumerate(headers):
        header_cell = vote_table.cell(0, index)
        _c(header_cell, header, size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(header_cell, "E8E8E8")

    for row_index, agenda in enumerate(agendas, 1):
        _c(vote_table.cell(row_index, 0), _render_value(str(agenda), variables), size=9)
        _c(vote_table.cell(row_index, 1), "", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _c(vote_table.cell(row_index, 2), "", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _p(doc, _render_value(vote_note, variables), size=9, space_before=6, space_after=16)

    _p(doc, assembly_date, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=12)

    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.columns[0].width = Cm(4.0)
    sign_table.columns[1].width = Cm(8.0)
    _c(sign_table.cell(0, 0), "조합원명", size=10, bold=True)
    _c(sign_table.cell(0, 1), "", size=10)
    _c(sign_table.cell(1, 0), "약정좌수", size=10, bold=True)
    _c(sign_table.cell(1, 1), "", size=10)
    _c(sign_table.cell(2, 0), "서명/날인", size=10, bold=True)
    _c(sign_table.cell(2, 1), "(인)", size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_table_borders(sign_table)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer