from __future__ import annotations

from datetime import date
from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from sqlalchemy.orm import Session

from models.investment import Investment, PortfolioCompany
from models.investment_review import InvestmentReview
from models.vote_record import VoteRecord

from .layout_utils import add_paragraph, create_base_document, set_cell_shading, set_cell_text, set_table_borders


def _fmt_date(value: date | None) -> str:
    if value is None:
        return "-"
    return value.isoformat()


def _fmt_money(value: float | int | None) -> str:
    if value is None:
        return "-"
    return f"{int(round(float(value))):,}원"


def _safe_text(value: str | None) -> str:
    return (value or "").strip() or "-"


def build_irc_minutes(investment_review_id: int, db: Session) -> bytes:
    """Build D-04 IRC minutes docx bytes."""
    review = db.get(InvestmentReview, investment_review_id)
    if not review:
        raise LookupError("투심위 심의 데이터를 찾을 수 없습니다.")

    investment = db.get(Investment, review.investment_id) if review.investment_id else None
    company = db.get(PortfolioCompany, investment.company_id) if investment else None

    vote_rows: list[VoteRecord] = []
    if review.investment_id:
        vote_rows = (
            db.query(VoteRecord)
            .filter(VoteRecord.investment_id == review.investment_id)
            .order_by(VoteRecord.date.asc(), VoteRecord.id.asc())
            .all()
        )
    if not vote_rows and company is not None:
        vote_rows = (
            db.query(VoteRecord)
            .filter(VoteRecord.company_id == company.id)
            .order_by(VoteRecord.date.asc(), VoteRecord.id.asc())
            .all()
        )

    approved = sum(1 for row in vote_rows if (row.decision or "").strip() in {"승인", "찬성", "가결"})
    rejected = sum(1 for row in vote_rows if (row.decision or "").strip() in {"반려", "반대", "부결"})

    doc = create_base_document()
    add_paragraph(
        doc,
        "투자심의위원회 의사록",
        size=16,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    info_rows = [
        ("회의일", _fmt_date(review.committee_date or review.decision_date)),
        ("대상기업", company.name if company else review.company_name),
        ("검토단계", _safe_text(review.status)),
        ("심사역", _safe_text(review.reviewer)),
        ("희망 투자금액", _fmt_money(review.target_amount)),
        ("투자수단", _safe_text(review.instrument)),
        ("의결 결과", _safe_text(review.decision_result)),
    ]

    info_table = doc.add_table(rows=len(info_rows), cols=2)
    info_table.columns[0].width = Cm(5.0)
    info_table.columns[1].width = Cm(11.0)
    set_table_borders(info_table)
    for index, (label, value) in enumerate(info_rows):
        set_cell_text(info_table.cell(index, 0), label, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(info_table.cell(index, 1), value, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        if index % 2 == 0:
            set_cell_shading(info_table.cell(index, 0), "F5F5F5")

    add_paragraph(doc, "", size=9, space_before=8)
    add_paragraph(doc, "안건", size=11, bold=True, space_after=2)
    add_paragraph(doc, _safe_text(review.review_opinion), size=10, space_after=6)
    add_paragraph(doc, "위원회 의견", size=11, bold=True, space_after=2)
    add_paragraph(doc, _safe_text(review.committee_opinion), size=10, space_after=8)

    add_paragraph(doc, f"표결 요약: 승인 {approved}건 / 반려 {rejected}건 / 총 {len(vote_rows)}건", size=10, space_after=4)

    vote_table = doc.add_table(rows=max(1, len(vote_rows)) + 1, cols=4)
    vote_table.columns[0].width = Cm(3.0)
    vote_table.columns[1].width = Cm(6.0)
    vote_table.columns[2].width = Cm(2.5)
    vote_table.columns[3].width = Cm(4.5)
    set_table_borders(vote_table)

    headers = ["일자", "안건", "결과", "메모"]
    for col, header in enumerate(headers):
        set_cell_text(vote_table.cell(0, col), header, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(vote_table.cell(0, col), "EFEFEF")

    if vote_rows:
        for row_index, row in enumerate(vote_rows, start=1):
            set_cell_text(vote_table.cell(row_index, 0), _fmt_date(row.date), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(vote_table.cell(row_index, 1), _safe_text(row.agenda), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(vote_table.cell(row_index, 2), _safe_text(row.decision), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(vote_table.cell(row_index, 3), _safe_text(row.memo), alignment=WD_ALIGN_PARAGRAPH.LEFT)
    else:
        set_cell_text(vote_table.cell(1, 0), "-", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(vote_table.cell(1, 1), "등록된 표결 기록이 없습니다.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(vote_table.cell(1, 2), "-", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(vote_table.cell(1, 3), "-", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, "", size=9, space_before=10)
    add_paragraph(doc, "참석위원 서명", size=10, bold=True, space_after=2)
    add_paragraph(doc, "위원장: ____________________", size=10, space_after=1)
    add_paragraph(doc, "위원 1: ____________________", size=10, space_after=1)
    add_paragraph(doc, "위원 2: ____________________", size=10, space_after=1)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
