from __future__ import annotations

from datetime import date, timedelta
from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from sqlalchemy.orm import Session

from models.fund import Fund
from models.investment import Investment, PortfolioCompany

from .layout_utils import add_paragraph, create_base_document, set_cell_shading, set_cell_text, set_table_borders

REQUIRED_DOCS = [
    "재무제표 (분기/반기/결산)",
    "사업자등록증 사본",
    "주주명부",
    "법인등기부등본",
    "4대보험 가입자명부",
    "대표자 신용정보조회동의서",
    "기타 변동사항 보고서",
]


def _quarter_start(year: int, quarter: int) -> date:
    month = ((quarter - 1) * 3) + 1
    return date(year, month, 1)


def _safe_text(value: str | None, default: str = "-") -> str:
    return (value or "").strip() or default


def build_doc_request_letter(
    fund_id: int,
    investment_id: int,
    quarter: int,
    year: int,
    db: Session,
) -> bytes:
    """Build D-06 document request letter docx bytes."""
    fund = db.get(Fund, fund_id)
    if not fund:
        raise LookupError("조합 데이터를 찾을 수 없습니다.")

    investment = db.get(Investment, investment_id)
    if not investment or investment.fund_id != fund_id:
        raise LookupError("투자 데이터를 찾을 수 없습니다.")

    company = db.get(PortfolioCompany, investment.company_id)
    if not company:
        raise LookupError("피투자사 데이터를 찾을 수 없습니다.")

    base_date = _quarter_start(year, quarter)
    deadline = base_date + timedelta(days=14)
    doc_no = f"{year}-Q{quarter}-{fund.id:03d}-{investment.id:04d}"

    doc = create_base_document()
    add_paragraph(
        doc,
        "피투자사 분기 서류 제출 요청 공문",
        size=16,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    header_rows = [
        ("문서번호", doc_no),
        ("수신", f"{company.name} 대표이사"),
        ("참조", "재무/관리 담당자"),
        ("발신", f"{fund.name} 업무집행조합원"),
        ("작성일", date.today().isoformat()),
        ("제목", f"{year}년 {quarter}분기 제출서류 요청의 건"),
    ]
    header_table = doc.add_table(rows=len(header_rows), cols=2)
    header_table.columns[0].width = Cm(4.0)
    header_table.columns[1].width = Cm(12.0)
    set_table_borders(header_table)
    for idx, (label, value) in enumerate(header_rows):
        set_cell_text(header_table.cell(idx, 0), label, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(header_table.cell(idx, 1), value, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        if idx % 2 == 0:
            set_cell_shading(header_table.cell(idx, 0), "F5F5F5")

    add_paragraph(doc, "", size=9, space_before=8)
    add_paragraph(
        doc,
        f"당 조합의 분기 점검 및 내부보고회 준비를 위해 아래 서류 제출을 요청드립니다.",
        size=10,
        space_after=4,
    )
    add_paragraph(
        doc,
        f"- 대상 기업: {company.name} ({_safe_text(company.business_number)})",
        size=10,
        space_before=1,
    )
    add_paragraph(
        doc,
        f"- 제출 기한: {deadline.isoformat()}",
        size=10,
        space_before=1,
        space_after=6,
    )

    docs_table = doc.add_table(rows=len(REQUIRED_DOCS) + 1, cols=2)
    docs_table.columns[0].width = Cm(1.8)
    docs_table.columns[1].width = Cm(14.2)
    set_table_borders(docs_table)
    set_cell_text(docs_table.cell(0, 0), "No", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(docs_table.cell(0, 1), "요청 서류", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_shading(docs_table.cell(0, 0), "EFEFEF")
    set_cell_shading(docs_table.cell(0, 1), "EFEFEF")
    for index, item in enumerate(REQUIRED_DOCS, start=1):
        set_cell_text(docs_table.cell(index, 0), str(index), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(docs_table.cell(index, 1), item, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_paragraph(doc, "", size=9, space_before=8)
    add_paragraph(
        doc,
        f"제출 방식: 이메일 회신 또는 V:ON ERP 업로드",
        size=10,
        space_after=1,
    )
    add_paragraph(doc, f"문의: {_safe_text(fund.fund_manager, '담당 운용역')}", size=10, space_after=10)
    add_paragraph(doc, date.today().isoformat(), size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)
    add_paragraph(
        doc,
        f"{fund.name} / 업무집행조합원 {_safe_text(fund.gp)} (인)",
        size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
