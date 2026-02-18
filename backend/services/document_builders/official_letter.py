from __future__ import annotations

from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from .layout_utils import (
    add_paragraph,
    add_run,
    create_base_document,
    resolve_layout_scale,
    set_cell_shading,
    set_cell_text,
    set_table_borders,
)

DEFAULT_ATTACHMENTS = [
    {"no": "1", "name": "조합규약(안)", "ref": "별첨1", "stamp_required": False},
    {"no": "2", "name": "조합규약(안)_별표3 조합원 동의서", "ref": "별표3", "stamp_required": True},
    {"no": "3", "name": "투자의사결정 심의기구 운영방안", "ref": "별첨2", "stamp_required": False},
]

DEFAULT_COVER_ATTACHMENTS = ["결성총회 소집통지서 1부", "결성총회 의안설명서 1부"]


def _as_dict(value) -> dict:
    return value if isinstance(value, dict) else {}


def _as_list(value) -> list:
    return value if isinstance(value, list) else []


def _render_value(text: str, variables: dict) -> str:
    rendered = text
    for key, value in variables.items():
        if key.startswith("__"):
            continue
        rendered = rendered.replace("{{" + key + "}}", str(value))
    return rendered


def build_official_letter(variables: dict) -> BytesIO:
    """Build the official letter document used in formation workflow."""
    custom = _as_dict(variables.get("__custom_data__"))
    company_header = _as_dict(custom.get("company_header"))
    payment_info = _as_dict(custom.get("payment_info"))

    layout_scale = resolve_layout_scale(custom)

    company_name = str(company_header.get("company_name") or variables.get("gp_name", "트리거투자파트너스(유)"))
    address = str(company_header.get("address") or variables.get("company_address", "서울특별시 마포구 양화로7길 70"))
    tel = str(company_header.get("tel") or variables.get("company_tel", "02-0000-0000"))
    fax = str(company_header.get("fax") or variables.get("company_fax", "02-0000-0000"))

    body_default = (
        "{{assembly_date}} 개최 예정인 {{fund_name}} 결성총회 관련 출자금 납입 및 제출서류를 안내드립니다. "
        "아래 내용을 확인해 주시기 바랍니다."
    )
    body_text = _render_value(str(custom.get("body_text") or body_default), variables)

    payment_note_default = "* 조합 규약에 따라 납입기한 이전 입금은 총회일 입금으로 간주될 수 있습니다."
    bank_account = str(payment_info.get("bank_account") or "(국민은행) 000-000000-00-000")
    payment_note = _render_value(str(payment_info.get("note") or payment_note_default), variables)

    attachments_raw = _as_list(custom.get("attachments")) or DEFAULT_ATTACHMENTS
    cover_attachments = _as_list(custom.get("cover_attachments")) or DEFAULT_COVER_ATTACHMENTS
    required_documents_text = _render_value(
        str(custom.get("required_documents_text") or "신분증 사본, 개인인감증명서"),
        variables,
    )

    attachments: list[tuple[str, str, str, str]] = []
    for index, item in enumerate(attachments_raw, start=1):
        row = item if isinstance(item, dict) else {}
        no = str(row.get("no") or index)
        name = _render_value(str(row.get("name") or ""), variables)
        ref = _render_value(str(row.get("ref") or ""), variables)
        stamp_mark = "O" if bool(row.get("stamp_required", False)) else ""
        attachments.append((no, name, ref, stamp_mark))

    doc = create_base_document(scale=layout_scale)

    def _p(*args, **kwargs):
        kwargs.setdefault("scale", layout_scale)
        return add_paragraph(*args, **kwargs)

    def _r(*args, **kwargs):
        kwargs.setdefault("scale", layout_scale)
        return add_run(*args, **kwargs)

    def _c(*args, **kwargs):
        kwargs.setdefault("scale", layout_scale)
        return set_cell_text(*args, **kwargs)

    title_para = _p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _r(title_para, company_name, size=17, bold=True)
    _r(title_para, "   T", size=17, bold=True, color=(0, 100, 180))

    contact_para = _p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    _r(contact_para, f"{address}   TEL {tel}   FAX {fax}", size=8, color=(120, 120, 120))

    _p(doc, "=" * 64, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    info_rows = [
        ("날 짜", str(variables.get("document_date", ""))),
        ("문서번호", str(variables.get("document_number", ""))),
        ("수 신", f"{variables.get('fund_name', '')} 조합원"),
        ("참 조", f"업무 담당자 {company_header.get('contact_name', '000')} / 이메일 {company_header.get('contact_email', '000')}") ,
        ("제 목", f"{variables.get('fund_name', '')} 출자금 납입 통지의 건"),
    ]

    info_table = doc.add_table(rows=len(info_rows), cols=2)
    info_table.columns[0].width = Cm(2.7)
    info_table.columns[1].width = Cm(13.1)
    for index, (label, value) in enumerate(info_rows):
        _c(info_table.cell(index, 0), label, size=10, bold=True)
        _c(info_table.cell(index, 1), _render_value(value, variables), size=10)

    _p(doc, space_before=10)
    _p(doc, body_text, size=10, space_after=6)

    assembly_date = str(variables.get("assembly_date", ""))
    assembly_time = str(variables.get("assembly_time", "오전 10시"))
    assembly_method = str(variables.get("assembly_method", "서면결의"))

    _p(doc, f"1) 결성총회 일시 : {assembly_date} {assembly_time}", size=10, space_before=2)
    _p(doc, f"2) 결성총회 방법 : {assembly_method}", size=10, space_before=2)
    _p(doc, "3) 출자이행통지", size=10, space_before=2, space_after=6)

    _p(doc, "- 다 음 -", size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=6)
    _p(doc, f"납입기한 : {assembly_date} {assembly_time}까지", size=10, space_before=2)
    _p(doc, f"납입계좌 : {_render_value(bank_account, variables)}", size=10, space_before=2)
    _p(doc, payment_note, size=9, space_before=2, space_after=10)

    _p(doc, "[첨부서류]", size=10, bold=True, space_before=6, space_after=2)
    for index, item in enumerate(cover_attachments, start=1):
        _p(doc, f"{index}. {_render_value(str(item), variables)}", size=10, space_before=1)

    _p(doc, "※ 결성총회 의안설명서 별첨 자료", size=9, bold=True, space_before=6, space_after=2)

    table = doc.add_table(rows=len(attachments) + 1, cols=4)
    set_table_borders(table)
    table.columns[0].width = Cm(1.2)
    table.columns[1].width = Cm(9.3)
    table.columns[2].width = Cm(2.4)
    table.columns[3].width = Cm(1.9)

    headers = ["No.", "목 록", "해당 자료", "날인"]
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        _c(cell, header, size=9, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "E8E8E8")

    for row, (no, name, ref, stamp) in enumerate(attachments, start=1):
        _c(table.cell(row, 0), no, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _c(table.cell(row, 1), name, size=9)
        _c(table.cell(row, 2), ref, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _c(table.cell(row, 3), stamp, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _p(doc, space_before=8)
    _p(doc, "[조합원 제출서류]", size=10, bold=True, space_after=2)
    _p(doc, required_documents_text, size=10, space_after=14)

    fund_name = str(variables.get("fund_name", ""))
    spaced_name = "  ".join(fund_name) if fund_name else ""
    sign_name = _p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=14)
    _r(sign_name, spaced_name, size=15, bold=True)

    sign_gp = _p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=3)
    _r(sign_gp, f"업무집행조합원 {variables.get('gp_name', '')}", size=11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer