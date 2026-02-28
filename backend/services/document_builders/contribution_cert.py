from __future__ import annotations

from datetime import date
from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from sqlalchemy.orm import Session

from models.fund import Fund, LP

from .layout_utils import add_paragraph, create_base_document, set_cell_text, set_table_borders


def _fmt_money(value: float | int | None) -> str:
    if value is None:
        return "-"
    return f"{int(round(float(value))):,}원"


def _fmt_percent(value: float | int | None) -> str:
    if value is None:
        return "-"
    return f"{float(value):.2f}%"


def build_contribution_cert(fund_id: int, lp_id: int, db: Session) -> bytes:
    """Build D-03 contribution certificate docx bytes."""
    fund = db.get(Fund, fund_id)
    if not fund:
        raise LookupError("조합 데이터를 찾을 수 없습니다.")

    lp = db.get(LP, lp_id)
    if not lp or lp.fund_id != fund_id:
        raise LookupError("LP 데이터를 찾을 수 없습니다.")

    commitment = float(lp.commitment or 0)
    paid_in = float(lp.paid_in or 0)
    denominator = float(fund.commitment_total or 0) or commitment
    ownership_pct = (paid_in / denominator * 100) if denominator > 0 else 0.0

    doc = create_base_document()
    add_paragraph(
        doc,
        "출자증서",
        size=18,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
    )

    add_paragraph(
        doc,
        f"본 증서는 아래 조합원의 출자금 납입 사실을 확인하기 위해 발급합니다.",
        size=10,
        space_after=8,
    )

    table = doc.add_table(rows=6, cols=2)
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(11.0)
    set_table_borders(table)

    rows = [
        ("조합명", fund.name),
        ("조합원(LP)", lp.name),
        ("약정금액", _fmt_money(lp.commitment)),
        ("납입금액", _fmt_money(lp.paid_in)),
        ("지분율", _fmt_percent(ownership_pct)),
        ("발급일", date.today().isoformat()),
    ]
    for index, (label, value) in enumerate(rows):
        set_cell_text(table.cell(index, 0), label, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.cell(index, 1), value, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_paragraph(doc, "", size=9, space_before=14)
    add_paragraph(doc, date.today().isoformat(), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_paragraph(doc, fund.name, size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(
        doc,
        f"업무집행조합원 {_safe_text(fund.gp)} (인)",
        size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _safe_text(value: str | None) -> str:
    return (value or "").strip() or "-"
