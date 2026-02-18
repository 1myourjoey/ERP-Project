from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

DEFAULT_FONT_NAME = "Malgun Gothic"


def _to_float(value, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def clamp_layout_scale(value: float) -> float:
    return max(0.74, min(1.0, _to_float(value, 1.0)))


def resolve_layout_scale(custom_data: dict | None) -> float:
    if not isinstance(custom_data, dict):
        return 1.0
    layout_fit = custom_data.get("__layout_fit__")
    if not isinstance(layout_fit, dict):
        return 1.0
    if not bool(layout_fit.get("enabled", False)):
        return 1.0
    return clamp_layout_scale(layout_fit.get("scale", 1.0))


def _scaled_size(size: int, scale: float, minimum: int = 7) -> int:
    return max(minimum, int(round(size * scale)))


def create_base_document(scale: float = 1.0) -> Document:
    """Create a base A4 document with balanced margins and optional fit scale."""
    fit_scale = clamp_layout_scale(scale)
    compression = 1.0 - fit_scale

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(max(1.4, 2.0 - (compression * 1.2)))
    section.bottom_margin = Cm(max(1.4, 2.0 - (compression * 1.2)))

    # Keep left/right margins symmetric so fit mode never looks unbalanced.
    side_margin = max(1.6, 2.5 - (compression * 2.0))
    section.left_margin = Cm(side_margin)
    section.right_margin = Cm(side_margin)
    return doc


def add_run(
    paragraph,
    text: str,
    size: int = 10,
    bold: bool = False,
    font_name: str = DEFAULT_FONT_NAME,
    color=None,
    scale: float = 1.0,
):
    fit_scale = clamp_layout_scale(scale)
    run = paragraph.add_run(text)
    run.font.size = Pt(_scaled_size(size, fit_scale))
    run.bold = bold
    run.font.name = font_name

    r_pr = run._element.get_or_add_rPr()
    r_fonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
    r_pr.append(r_fonts)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_paragraph(
    doc: Document,
    text: str = "",
    size: int = 10,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    space_before: int = 0,
    space_after: int = 0,
    font_name: str = DEFAULT_FONT_NAME,
    scale: float = 1.0,
):
    fit_scale = clamp_layout_scale(scale)
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_before = Pt(max(0, int(round(space_before * fit_scale))))
    para.paragraph_format.space_after = Pt(max(0, int(round(space_after * fit_scale))))
    para.paragraph_format.line_spacing = max(1.1, 1.45 - ((1.0 - fit_scale) * 0.8))
    if text:
        add_run(para, text, size=size, bold=bold, font_name=font_name, scale=fit_scale)
    return para


def set_cell_text(
    cell,
    text: str,
    size: int = 9,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    font_name: str = DEFAULT_FONT_NAME,
    scale: float = 1.0,
):
    fit_scale = clamp_layout_scale(scale)
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = alignment
    para.paragraph_format.line_spacing = max(1.08, 1.35 - ((1.0 - fit_scale) * 0.7))
    add_run(para, text, size=size, bold=bold, font_name=font_name, scale=fit_scale)


def set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, size: int = 4, color: str = "000000"):
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    table._tbl.tblPr.append(borders)