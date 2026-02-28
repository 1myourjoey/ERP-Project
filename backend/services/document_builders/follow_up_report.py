from __future__ import annotations

from datetime import date
from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from sqlalchemy.orm import Session

from models.internal_review import CompanyReview, InternalReview
from models.investment import Investment, PortfolioCompany

from .layout_utils import (
    add_paragraph,
    create_base_document,
    set_cell_shading,
    set_cell_text,
    set_table_borders,
)


def _fmt_date(value: date | None) -> str:
    if value is None:
        return "-"
    return value.isoformat()


def _fmt_money(value: float | int | None) -> str:
    if value is None:
        return "-"
    return f"{int(round(float(value))):,}원"


def _fmt_percent(value: float | int | None) -> str:
    if value is None:
        return "-"
    return f"{float(value):.2f}%"


def _fmt_signed(value: int | None) -> str:
    if value is None:
        return "-"
    if value > 0:
        return f"+{value}"
    return str(value)


def _safe_text(value: str | None) -> str:
    return (value or "").strip() or "-"


def build_follow_up_report(company_review_id: int, db: Session) -> bytes:
    """Build D-01 follow-up report docx bytes."""
    company_review = db.get(CompanyReview, company_review_id)
    if not company_review:
        raise LookupError("회사별 리뷰 데이터를 찾을 수 없습니다.")

    review = db.get(InternalReview, company_review.review_id)
    if not review:
        raise LookupError("내부보고회 데이터를 찾을 수 없습니다.")

    investment = db.get(Investment, company_review.investment_id)
    if not investment:
        raise LookupError("투자 데이터를 찾을 수 없습니다.")

    company = db.get(PortfolioCompany, investment.company_id)
    if not company:
        raise LookupError("피투자사 데이터를 찾을 수 없습니다.")

    doc = create_base_document()
    add_paragraph(
        doc,
        "후속관리보고서",
        size=16,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    rows = [
        ("기준일", _fmt_date(review.reference_date)),
        ("작성일", _fmt_date(date.today())),
        ("기업명", company.name),
        ("대표자", _safe_text(company.ceo)),
        ("사업자번호", _safe_text(company.business_number)),
        ("업종", _safe_text(company.industry)),
        ("투자일", _fmt_date(investment.investment_date)),
        ("투자금액", _fmt_money(investment.amount)),
        ("지분율", _fmt_percent(investment.ownership_pct)),
        ("분기 매출액", _fmt_money(company_review.quarterly_revenue)),
        ("분기 영업이익", _fmt_money(company_review.quarterly_operating_income)),
        ("분기 당기순이익", _fmt_money(company_review.quarterly_net_income)),
        ("총자산", _fmt_money(company_review.total_assets)),
        ("총부채", _fmt_money(company_review.total_liabilities)),
        ("현금성자산", _fmt_money(company_review.cash_and_equivalents)),
        ("임직원 수", str(company_review.employee_count or "-")),
        ("인원 변동", _fmt_signed(company_review.employee_change)),
        ("자산건전성 등급", _safe_text(company_review.asset_rating)),
        ("손상차손", _safe_text(company_review.impairment_type)),
        ("이사회 참석", _safe_text(company_review.board_attendance)),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(11.0)
    set_table_borders(table)

    for row_index, (label, value) in enumerate(rows):
        header_cell = table.cell(row_index, 0)
        value_cell = table.cell(row_index, 1)
        set_cell_text(
            header_cell,
            label,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        set_cell_text(value_cell, value, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        if row_index % 2 == 0:
            set_cell_shading(header_cell, "F5F5F5")

    add_paragraph(doc, "", size=9, space_before=6)
    note_table = doc.add_table(rows=3, cols=2)
    note_table.columns[0].width = Cm(5.0)
    note_table.columns[1].width = Cm(11.0)
    set_table_borders(note_table)

    note_rows = [
        ("주요 이슈", _safe_text(company_review.key_issues)),
        ("후속 조치", _safe_text(company_review.follow_up_actions)),
        ("종합 의견", _safe_text(company_review.investment_opinion)),
    ]
    for idx, (label, value) in enumerate(note_rows):
        header_cell = note_table.cell(idx, 0)
        value_cell = note_table.cell(idx, 1)
        set_cell_text(header_cell, label, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(value_cell, value, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(header_cell, "F5F5F5")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
