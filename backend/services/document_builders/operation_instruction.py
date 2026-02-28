from __future__ import annotations

from datetime import date
from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from sqlalchemy.orm import Session

from models.fund import Fund
from models.transaction import Transaction

from .layout_utils import (
    add_paragraph,
    create_base_document,
    set_cell_shading,
    set_cell_text,
    set_table_borders,
    to_korean_currency,
)

INSTRUCTION_TYPE_LABEL: dict[str, str] = {
    "investment": "투자집행",
    "follow_on": "후속투자",
    "dividend": "배분금 지급",
    "exit": "회수금 수령",
    "conversion": "전환거래",
    "other": "기타 운용지시",
}


def _safe_text(value: str | None, default: str = "-") -> str:
    return (value or "").strip() or default


def _resolve_to_account(transaction: Transaction) -> str:
    if transaction.counterparty:
        return transaction.counterparty.strip()
    if transaction.type == "investment":
        return "피투자사 지정 계좌"
    if transaction.type == "dividend":
        return "LP 분배계좌"
    if transaction.type == "exit":
        return "조합 수납계좌"
    return "상대방 계좌"


def _resolve_to_account_holder(transaction: Transaction) -> str:
    if transaction.counterparty:
        return transaction.counterparty.strip()
    if transaction.type == "investment":
        return "피투자사"
    if transaction.type == "dividend":
        return "LP"
    if transaction.type == "exit":
        return "조합"
    return "-"


def build_operation_instruction(transaction_id: int, db: Session) -> bytes:
    """Build D-02 operation instruction docx bytes."""
    transaction = db.get(Transaction, transaction_id)
    if not transaction:
        raise LookupError("거래 데이터를 찾을 수 없습니다.")

    fund = db.get(Fund, transaction.fund_id)
    if not fund:
        raise LookupError("조합 데이터를 찾을 수 없습니다.")

    amount = float(transaction.amount or 0)
    amount_text = f"{int(round(amount)):,}원"
    amount_korean = to_korean_currency(amount)

    doc = create_base_document()
    add_paragraph(
        doc,
        "운용지시서",
        size=16,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    rows = [
        ("조합명", fund.name),
        ("지시유형", INSTRUCTION_TYPE_LABEL.get(transaction.type, transaction.type or "기타")),
        ("지시일", date.today().isoformat()),
        ("거래일", transaction.transaction_date.isoformat() if transaction.transaction_date else "-"),
        ("금액", amount_text),
        ("한글금액", amount_korean),
        ("출금계좌", _safe_text(fund.account_number)),
        ("입금계좌", _resolve_to_account(transaction)),
        ("예금주", _resolve_to_account_holder(transaction)),
        ("목적", _safe_text(transaction.memo)),
        ("수탁사", _safe_text(fund.trustee)),
        ("업무집행조합원", _safe_text(fund.gp)),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(11.0)
    set_table_borders(table)

    for index, (label, value) in enumerate(rows):
        header_cell = table.cell(index, 0)
        value_cell = table.cell(index, 1)
        set_cell_text(header_cell, label, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(value_cell, value, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        if index % 2 == 0:
            set_cell_shading(header_cell, "F5F5F5")

    add_paragraph(doc, "", size=9, space_before=8)
    add_paragraph(doc, "상기 내용과 같이 자금 이체를 요청드립니다.", size=10, space_after=16)
    add_paragraph(doc, date.today().isoformat(), size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)
    add_paragraph(
        doc,
        f"{_safe_text(fund.gp, '업무집행조합원')} (인)",
        size=11,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
