from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from sqlalchemy import and_, func
from sqlalchemy.orm import Session

from models.fee import ManagementFee, PerformanceFeeSimulation
from models.fund import Fund, LP
from models.investment import Investment, PortfolioCompany
from models.phase3 import CapitalCall, Distribution, DistributionDetail, ExitTrade
from models.valuation import Valuation
from services.performance_calculator import calculate_fund_performance


def _quarter_bounds(year: int, quarter: int) -> tuple[date, date]:
    if quarter not in {1, 2, 3, 4}:
        raise ValueError("quarter must be between 1 and 4")
    start_month = ((quarter - 1) * 3) + 1
    start = date(year, start_month, 1)
    if quarter == 4:
        end = date(year, 12, 31)
    else:
        end = date(year, start_month + 3, 1) - date.resolution
    return start, end


def _latest_valuation_map(db: Session, fund_id: int) -> dict[int, Valuation]:
    rows = (
        db.query(Valuation)
        .filter(Valuation.fund_id == fund_id)
        .order_by(Valuation.investment_id.asc(), Valuation.as_of_date.desc(), Valuation.id.desc())
        .all()
    )
    mapping: dict[int, Valuation] = {}
    for row in rows:
        if row.investment_id in mapping:
            continue
        mapping[row.investment_id] = row
    return mapping


async def collect_lp_report_data(
    db: Session,
    fund_id: int,
    year: int,
    quarter: int,
) -> dict:
    fund = db.get(Fund, fund_id)
    if not fund:
        raise ValueError("fund not found")

    period_start, period_end = _quarter_bounds(year, quarter)

    investments = (
        db.query(Investment)
        .filter(Investment.fund_id == fund_id)
        .order_by(Investment.investment_date.asc().nullslast(), Investment.id.asc())
        .all()
    )
    latest_valuations = _latest_valuation_map(db, fund_id)

    portfolio: list[dict] = []
    for investment in investments:
        company = db.get(PortfolioCompany, investment.company_id)
        latest = latest_valuations.get(investment.id)
        current_valuation = float(
            latest.total_fair_value
            if latest and latest.total_fair_value is not None
            else latest.value
            if latest
            else 0
        )
        invested_amount = float(investment.amount or 0)
        unrealized_gain = current_valuation - invested_amount
        portfolio.append(
            {
                "company": company.name if company else f"Investment #{investment.id}",
                "investment_date": investment.investment_date.isoformat() if investment.investment_date else None,
                "amount": invested_amount,
                "current_valuation": round(current_valuation, 2),
                "unrealized_gain": round(unrealized_gain, 2),
                "status": investment.status,
            }
        )

    performance = await calculate_fund_performance(db, fund_id, as_of_date=period_end)

    mgmt_fee_ytd = float(
        db.query(func.coalesce(func.sum(ManagementFee.fee_amount), 0))
        .filter(ManagementFee.fund_id == fund_id, ManagementFee.year == year)
        .scalar()
        or 0
    )
    perf_sim = (
        db.query(PerformanceFeeSimulation)
        .filter(PerformanceFeeSimulation.fund_id == fund_id)
        .order_by(PerformanceFeeSimulation.simulation_date.desc(), PerformanceFeeSimulation.id.desc())
        .first()
    )
    performance_fee_ytd = float(perf_sim.carry_amount or 0) if perf_sim else 0.0

    total_commitment = float(sum(float(lp.commitment or 0) for lp in fund.lps))
    total_paid_in = float(sum(float(lp.paid_in or 0) for lp in fund.lps))
    total_distributed = float(
        db.query(func.coalesce(func.sum(Distribution.principal_total + Distribution.profit_total), 0))
        .filter(Distribution.fund_id == fund_id)
        .scalar()
        or 0
    )

    events: list[dict] = []

    quarter_investments = [
        row for row in investments
        if row.investment_date is not None and period_start <= row.investment_date <= period_end
    ]
    for row in quarter_investments:
        company = db.get(PortfolioCompany, row.company_id)
        events.append(
            {
                "date": row.investment_date.isoformat(),
                "type": "investment",
                "description": f"신규 투자: {company.name if company else row.id}",
                "amount": float(row.amount or 0),
            }
        )

    exit_rows = (
        db.query(ExitTrade)
        .filter(ExitTrade.fund_id == fund_id, ExitTrade.trade_date >= period_start, ExitTrade.trade_date <= period_end)
        .all()
    )
    for row in exit_rows:
        company = db.get(PortfolioCompany, row.company_id)
        events.append(
            {
                "date": row.trade_date.isoformat(),
                "type": "exit",
                "description": f"엑시트: {company.name if company else row.id}",
                "amount": float(row.net_amount if row.net_amount is not None else row.amount or 0),
            }
        )

    capital_calls = (
        db.query(CapitalCall)
        .filter(CapitalCall.fund_id == fund_id, CapitalCall.call_date >= period_start, CapitalCall.call_date <= period_end)
        .all()
    )
    for row in capital_calls:
        events.append(
            {
                "date": row.call_date.isoformat(),
                "type": "capital_call",
                "description": f"자본금 콜 #{row.id}",
                "amount": float(row.total_amount or 0),
            }
        )

    quarter_distributions = (
        db.query(Distribution)
        .filter(Distribution.fund_id == fund_id, Distribution.dist_date >= period_start, Distribution.dist_date <= period_end)
        .all()
    )
    for row in quarter_distributions:
        events.append(
            {
                "date": row.dist_date.isoformat(),
                "type": "distribution",
                "description": f"배분 #{row.id}",
                "amount": float(row.principal_total or 0) + float(row.profit_total or 0),
            }
        )

    events.sort(key=lambda item: item["date"])

    distribution_by_lp = dict(
        db.query(
            DistributionDetail.lp_id,
            func.coalesce(func.sum(DistributionDetail.distribution_amount), 0),
        )
        .join(Distribution, Distribution.id == DistributionDetail.distribution_id)
        .filter(Distribution.fund_id == fund_id)
        .group_by(DistributionDetail.lp_id)
        .all()
    )

    lp_summary: list[dict] = []
    for lp in sorted(fund.lps, key=lambda row: row.id):
        commitment = float(lp.commitment or 0)
        paid_in = float(lp.paid_in or 0)
        distributions = float(distribution_by_lp.get(lp.id, 0) or 0)
        nav_share = (
            (paid_in / total_paid_in) * float(performance.get("residual_value") or 0)
            if total_paid_in > 0
            else 0
        )
        lp_summary.append(
            {
                "lp_name": lp.name,
                "commitment": commitment,
                "paid_in": paid_in,
                "distributions": distributions,
                "nav_share": round(nav_share, 2),
            }
        )

    return {
        "fund": {
            "id": fund.id,
            "name": fund.name,
            "type": fund.type,
            "formation_date": fund.formation_date.isoformat() if fund.formation_date else None,
            "status": fund.status,
            "total_commitment": total_commitment,
            "total_paid_in": total_paid_in,
            "total_distributed": total_distributed,
        },
        "period": {
            "year": year,
            "quarter": quarter,
            "start_date": period_start.isoformat(),
            "end_date": period_end.isoformat(),
        },
        "portfolio": portfolio,
        "performance": performance,
        "fees": {
            "mgmt_fee_ytd": round(mgmt_fee_ytd, 2),
            "performance_fee_ytd": round(performance_fee_ytd, 2),
            "total_fees_ytd": round(mgmt_fee_ytd + performance_fee_ytd, 2),
        },
        "capital": {
            "total_commitment": total_commitment,
            "total_paid_in": total_paid_in,
            "total_distributed": total_distributed,
            "remaining_commitment": max(total_commitment - total_paid_in, 0.0),
            "paid_in_ratio": round((total_paid_in / total_commitment) * 100, 2) if total_commitment > 0 else 0.0,
        },
        "events": events,
        "lp_summary": lp_summary,
    }


async def generate_lp_report_docx(
    db: Session,
    fund_id: int,
    year: int,
    quarter: int,
) -> bytes:
    payload = await collect_lp_report_data(db, fund_id, year, quarter)

    doc = Document()
    doc.add_heading(f"{payload['fund']['name']} LP 분기 보고서", level=1)
    doc.add_paragraph(f"보고기간: {payload['period']['year']}년 {payload['period']['quarter']}분기")
    doc.add_paragraph(f"상태: {payload['fund']['status']} / 유형: {payload['fund']['type']}")

    doc.add_heading("1. 펀드 개요", level=2)
    doc.add_paragraph(f"총약정: {payload['fund']['total_commitment']:,.0f}")
    doc.add_paragraph(f"누적 납입: {payload['fund']['total_paid_in']:,.0f}")
    doc.add_paragraph(f"누적 배분: {payload['fund']['total_distributed']:,.0f}")

    doc.add_heading("2. 수익률", level=2)
    perf = payload["performance"]
    irr_text = f"{perf['irr'] * 100:.2f}%" if perf.get("irr") is not None else "N/A"
    doc.add_paragraph(f"IRR: {irr_text}")
    doc.add_paragraph(f"TVPI: {perf['tvpi']:.2f}x")
    doc.add_paragraph(f"DPI: {perf['dpi']:.2f}x")

    doc.add_heading("3. 포트폴리오", level=2)
    if payload["portfolio"]:
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "기업"
        hdr[1].text = "투자일"
        hdr[2].text = "투자금"
        hdr[3].text = "현재가치"
        hdr[4].text = "미실현손익"
        for row in payload["portfolio"]:
            cells = table.add_row().cells
            cells[0].text = str(row["company"])
            cells[1].text = row["investment_date"] or "-"
            cells[2].text = f"{float(row['amount']):,.0f}"
            cells[3].text = f"{float(row['current_valuation']):,.0f}"
            cells[4].text = f"{float(row['unrealized_gain']):,.0f}"
    else:
        doc.add_paragraph("해당 기간 포트폴리오 데이터가 없습니다.")

    doc.add_heading("4. 주요 이벤트", level=2)
    if payload["events"]:
        for event in payload["events"]:
            doc.add_paragraph(
                f"- {event['date']} | {event['type']} | {event['description']} | {float(event['amount']):,.0f}"
            )
    else:
        doc.add_paragraph("해당 기간 이벤트가 없습니다.")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
