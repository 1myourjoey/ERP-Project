from __future__ import annotations

import json
import os
import re
from datetime import date, datetime
from functools import lru_cache
from io import BytesIO

from models.document_template import DocumentTemplate
from models.fund import Fund, LP

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SPREADSHEET_MEDIA_TYPES = {
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xlsm": "application/vnd.ms-excel.sheet.macroEnabled.12",
}
MARKER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _to_datetime(value) -> datetime | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value
    if isinstance(value, date):
        return datetime.combine(value, datetime.min.time())
    try:
        return datetime.fromisoformat(str(value))
    except ValueError:
        return None


def replace_text_in_paragraph(paragraph, variables: dict):
    """Replace {{key}} placeholders in a paragraph."""
    full_text = paragraph.text
    for key, value in variables.items():
        placeholder = "{{" + key + "}}"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))
    if full_text != paragraph.text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full_text


def replace_text_in_table(table, variables: dict):
    """Replace placeholders in every table cell paragraph."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, variables)


def build_variables_for_fund(fund: Fund, lps: list[LP], extra: dict | None = None) -> dict:
    """Build a variable map from fund and LP data."""
    now = datetime.now()

    def fmt_date(value, include_day_name: bool = True) -> str:
        dt = _to_datetime(value)
        if not dt:
            return "미정"
        day_names = ["월요일", "화요일", "수요일", "목요일", "금요일", "토요일", "일요일"]
        if include_day_name:
            return f"{dt.year}년 {dt.month}월 {dt.day}일({day_names[dt.weekday()]})"
        return f"{dt.year}년 {dt.month}월 {dt.day}일"

    def fmt_date_short(value) -> str:
        dt = _to_datetime(value)
        if not dt:
            return "미정"
        return f"{dt.year}. {dt.month:02d}. {dt.day:02d}"

    def fmt_amount(value) -> str:
        if value in (None, ""):
            return "0"
        try:
            return f"{int(float(value)):,}"
        except (TypeError, ValueError):
            return "0"

    lp_lines: list[str] = []
    for index, lp in enumerate(lps, 1):
        commitment = getattr(lp, "commitment", None)
        lp_lines.append(f"  {index}. {lp.name} ({lp.type or ''}): 약정 {fmt_amount(commitment)}원")
    lp_list_text = "\n".join(lp_lines)

    total_commitment = sum(float(getattr(lp, "commitment", 0) or 0) for lp in lps)
    lp_count = len(lps)

    variables = {
        "fund_name": fund.name or "",
        "fund_type": fund.type or "",
        "fund_status": fund.status or "",
        "gp_name": fund.gp or "트리거투자파트너스(유)",
        "registration_number": fund.registration_number or "",
        "registration_date": fmt_date(fund.registration_date),
        "registration_date_short": fmt_date_short(fund.registration_date),
        "co_gp_name": fund.co_gp or "",
        "trustee": fund.trustee or "",
        "commitment_total": fmt_amount(fund.commitment_total),
        "commitment_total_raw": str(fund.commitment_total or 0),
        "formation_date": fmt_date(fund.formation_date),
        "formation_date_short": fmt_date_short(fund.formation_date),
        "today_date": fmt_date(now),
        "today_date_short": fmt_date_short(now),
        "lp_count": str(lp_count),
        "lp_list": lp_list_text,
        "total_commitment_amount": fmt_amount(total_commitment),
        "document_date": fmt_date_short(now),
    }

    if extra:
        variables.update(extra)
    return variables


def _resolve_template_path(file_path: str) -> str:
    if os.path.isabs(file_path):
        return file_path
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    return os.path.join(project_root, file_path)


def template_output_extension(template: DocumentTemplate) -> str:
    suffix = os.path.splitext(template.file_path or "")[1].lower()
    if suffix in SPREADSHEET_MEDIA_TYPES:
        return suffix
    return ".docx"


def template_output_media_type(template: DocumentTemplate) -> str:
    suffix = template_output_extension(template)
    return SPREADSHEET_MEDIA_TYPES.get(suffix, DOCX_MEDIA_TYPE)


def _replace_markers_in_text(text: str, variables: dict) -> str:
    def replacer(match: re.Match[str]) -> str:
        key = match.group(1).strip()
        if key not in variables:
            return match.group(0)
        return str(variables[key])

    return MARKER_PATTERN.sub(replacer, text)


def generate_document(template: DocumentTemplate, variables: dict) -> BytesIO:
    """Generate a .docx by replacing placeholders in a file-based template."""
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError("python-docx가 설치되어 있지 않습니다.") from exc

    template_path = _resolve_template_path(template.file_path or "")
    if not template.file_path or not os.path.exists(template_path):
        raise FileNotFoundError(f"템플릿 파일을 찾을 수 없습니다: {template_path}")

    if template_output_extension(template) != ".docx":
        raise ValueError(f"Unsupported Word template type: {template.file_path}")

    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, variables)

    for table in doc.tables:
        replace_text_in_table(table, variables)

    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            for paragraph in header.paragraphs:
                replace_text_in_paragraph(paragraph, variables)
        for footer in [section.footer, section.first_page_footer]:
            for paragraph in footer.paragraphs:
                replace_text_in_paragraph(paragraph, variables)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_spreadsheet(template: DocumentTemplate, variables: dict) -> BytesIO:
    try:
        from openpyxl import load_workbook
    except ModuleNotFoundError as exc:
        raise RuntimeError("openpyxl is not installed.") from exc

    template_path = _resolve_template_path(template.file_path or "")
    if not template.file_path or not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    suffix = template_output_extension(template)
    if suffix not in SPREADSHEET_MEDIA_TYPES:
        raise ValueError(f"Unsupported spreadsheet template type: {template.file_path}")

    workbook = load_workbook(
        template_path,
        keep_vba=(suffix == ".xlsm"),
        data_only=False,
    )
    try:
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        cell.value = _replace_markers_in_text(cell.value, variables)

            for header_footer_name in (
                "oddHeader",
                "evenHeader",
                "firstHeader",
                "oddFooter",
                "evenFooter",
                "firstFooter",
            ):
                header_footer = getattr(worksheet, header_footer_name, None)
                if header_footer is None:
                    continue
                for side_name in ("left", "center", "right"):
                    side = getattr(header_footer, side_name, None)
                    text = getattr(side, "text", None)
                    if isinstance(text, str) and text:
                        side.text = _replace_markers_in_text(text, variables)

        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        return buffer
    finally:
        workbook.close()


def extract_template_markers(template: DocumentTemplate) -> list[str]:
    suffix = template_output_extension(template)
    template_path = _resolve_template_path(template.file_path or "")
    if not template.file_path or not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template.file_path}")

    markers: set[str] = set()

    def scan_text(text: str) -> None:
        for match in MARKER_PATTERN.finditer(text):
            markers.add(match.group(1))

    if suffix == ".docx":
        from docx import Document

        document = Document(template_path)
        for paragraph in document.paragraphs:
            scan_text("".join(run.text for run in paragraph.runs))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        scan_text("".join(run.text for run in paragraph.runs))

        for section in document.sections:
            for header in (section.header, section.first_page_header):
                if header:
                    for paragraph in header.paragraphs:
                        scan_text("".join(run.text for run in paragraph.runs))
            for footer in (section.footer, section.first_page_footer):
                if footer:
                    for paragraph in footer.paragraphs:
                        scan_text("".join(run.text for run in paragraph.runs))
        return sorted(markers)

    if suffix not in SPREADSHEET_MEDIA_TYPES:
        raise ValueError(f"Unsupported template type: {template.file_path}")

    from openpyxl import load_workbook

    workbook = load_workbook(
        template_path,
        read_only=True,
        keep_vba=(suffix == ".xlsm"),
        data_only=False,
    )
    try:
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        scan_text(cell.value)
    finally:
        workbook.close()

    return sorted(markers)


@lru_cache(maxsize=1)
def _load_document_builders():
    try:
        from services.document_builders.assembly_notice import build_assembly_notice
        from services.document_builders.official_letter import build_official_letter
        from services.document_builders.written_resolution import build_written_resolution
    except ModuleNotFoundError as exc:
        raise RuntimeError("python-docx가 설치되어 있지 않습니다.") from exc

    return {
        "공문_결성총회_출자이행통지": build_official_letter,
        "첨부1_결성총회_소집통지서": build_assembly_notice,
        "별첨6_서면결의서": build_written_resolution,
    }


def generate_document_v2(template_name: str, variables: dict) -> BytesIO:
    """Generate a code-built .docx from a registered builder name."""
    builder = _load_document_builders().get(template_name)
    if not builder:
        raise ValueError(f"지원하지 않는 문서 유형: {template_name}")
    return builder(variables)


def _select_builder_for_template(template: DocumentTemplate, variables: dict):
    """Pick code builder by explicit name first, then by custom_data shape fallback."""
    builders = _load_document_builders()

    for key in [template.builder_name or "", template.name or ""]:
        if key in builders:
            return builders[key]

    custom_data = variables.get("__custom_data__")
    if isinstance(custom_data, dict):
        keys = set(custom_data.keys())
        if {"company_header", "payment_info", "attachments"} & keys:
            return builders.get("공문_결성총회_출자이행통지")
        if {"greeting", "regulation_article"} & keys:
            return builders.get("첨부1_결성총회_소집통지서")
        if {"introduction_text", "vote_note"} & keys:
            return builders.get("별첨6_서면결의서")

    return None


def generate_document_for_template(template: DocumentTemplate, variables: dict) -> BytesIO:
    """Use code builder first, then fallback to file-based replacement."""
    if template.custom_data and template.custom_data != "{}":
        try:
            custom = json.loads(template.custom_data)
            if isinstance(custom, dict) and "__custom_data__" not in variables:
                variables["__custom_data__"] = custom
        except json.JSONDecodeError:
            pass

    if template_output_extension(template) in SPREADSHEET_MEDIA_TYPES:
        return generate_spreadsheet(template, variables)

    builder = _select_builder_for_template(template, variables)
    if builder:
        return builder(variables)

    return generate_document(template, variables)
