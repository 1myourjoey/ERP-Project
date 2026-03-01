from __future__ import annotations

import re
from io import BytesIO
from zipfile import ZipFile

from docx import Document


class TextExtractor:
    """Extract plain text from DOCX/HWP/HWPX templates."""

    def extract_from_docx(self, file_bytes: bytes) -> str:
        document = Document(BytesIO(file_bytes))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))

        for section in document.sections:
            for header in (section.header, section.first_page_header):
                if not header:
                    continue
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text)
            for footer in (section.footer, section.first_page_footer):
                if not footer:
                    continue
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text)

        return "\n".join(parts).strip()

    def extract_from_hwp(self, file_bytes: bytes) -> str:
        try:
            import olefile  # type: ignore
        except Exception:
            return ""

        try:
            hwp = olefile.OleFileIO(BytesIO(file_bytes))
        except Exception:
            return ""

        try:
            if hwp.exists("PrvText"):
                text_bytes = hwp.openstream("PrvText").read()
                return text_bytes.decode("utf-16-le", errors="ignore").strip()

            sections = [row for row in hwp.listdir() if row and row[0] == "BodyText"]
            chunks: list[str] = []
            for section in sorted(sections):
                raw = hwp.openstream("/".join(section)).read()
                decoded = raw.decode("utf-16-le", errors="ignore").strip()
                if decoded:
                    chunks.append(decoded)
            return "\n".join(chunks).strip()
        except Exception:
            return ""
        finally:
            try:
                hwp.close()
            except Exception:
                pass

    def extract_from_hwpx(self, file_bytes: bytes) -> str:
        try:
            archive = ZipFile(BytesIO(file_bytes))
        except Exception:
            return ""

        text_chunks: list[str] = []
        with archive:
            candidates = [name for name in archive.namelist() if name.lower().endswith(".xml")]
            for name in candidates:
                try:
                    content = archive.read(name).decode("utf-8", errors="ignore")
                except Exception:
                    continue
                cleaned = re.sub(r"<[^>]+>", " ", content)
                cleaned = re.sub(r"\s+", " ", cleaned).strip()
                if cleaned:
                    text_chunks.append(cleaned)
        return "\n".join(text_chunks).strip()

    def extract(self, file_bytes: bytes, filename: str) -> str:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext == "docx":
            return self.extract_from_docx(file_bytes)
        if ext == "hwp":
            return self.extract_from_hwp(file_bytes)
        if ext == "hwpx":
            return self.extract_from_hwpx(file_bytes)
        raise ValueError(f"Unsupported template type: {ext or 'unknown'}")
