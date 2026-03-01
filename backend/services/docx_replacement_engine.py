from __future__ import annotations

import re
from io import BytesIO

from docx import Document


class DocxReplacementEngine:
    """Format-preserving DOCX marker replacement engine.

    Replaces ``{{marker}}`` values while preserving run-level style information.
    Markers split across multiple runs are also handled.
    """

    MARKER_PATTERN = re.compile(r"\{\{([^{}\s]+)\}\}")

    def replace(self, template_bytes: bytes, variables: dict[str, str]) -> bytes:
        document = Document(BytesIO(template_bytes))

        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, variables)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, variables)

        for section in document.sections:
            for header in (section.header, section.first_page_header):
                if header:
                    for paragraph in header.paragraphs:
                        self._replace_in_paragraph(paragraph, variables)
            for footer in (section.footer, section.first_page_footer):
                if footer:
                    for paragraph in footer.paragraphs:
                        self._replace_in_paragraph(paragraph, variables)

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def extract_markers(self, template_bytes: bytes) -> list[str]:
        document = Document(BytesIO(template_bytes))
        markers: set[str] = set()

        def scan_text(text: str) -> None:
            for match in self.MARKER_PATTERN.finditer(text):
                markers.add(match.group(1))

        for paragraph in document.paragraphs:
            scan_text("".join(run.text for run in paragraph.runs))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        scan_text("".join(run.text for run in paragraph.runs))

        for section in document.sections:
            for header in (section.header, section.first_page_header):
                if header:
                    for paragraph in header.paragraphs:
                        scan_text("".join(run.text for run in paragraph.runs))
            for footer in (section.footer, section.first_page_footer):
                if footer:
                    for paragraph in footer.paragraphs:
                        scan_text("".join(run.text for run in paragraph.runs))

        return sorted(markers)

    def _replace_in_paragraph(self, paragraph, variables: dict[str, str]) -> None:
        runs = list(paragraph.runs)
        if not runs:
            return

        full_text = "".join(run.text for run in runs)
        if "{{" not in full_text:
            return

        matches = list(self.MARKER_PATTERN.finditer(full_text))
        if not matches:
            return

        run_ranges = self._build_run_ranges(runs)

        for match in reversed(matches):
            marker_name = match.group(1)
            if marker_name not in variables:
                continue
            replacement = str(variables.get(marker_name, ""))
            self._replace_range(runs, run_ranges, match.start(), match.end(), replacement)

    @staticmethod
    def _build_run_ranges(runs) -> list[tuple[int, int, int]]:
        ranges: list[tuple[int, int, int]] = []
        cursor = 0
        for index, run in enumerate(runs):
            text = run.text or ""
            start = cursor
            end = cursor + len(text)
            ranges.append((index, start, end))
            cursor = end
        return ranges

    @staticmethod
    def _replace_range(
        runs,
        run_ranges: list[tuple[int, int, int]],
        start_pos: int,
        end_pos: int,
        replacement: str,
    ) -> None:
        affected = [row for row in run_ranges if row[1] < end_pos and row[2] > start_pos]
        if not affected:
            return

        first_index, first_start, _first_end = affected[0]
        last_index, last_start, _last_end = affected[-1]

        first_run = runs[first_index]
        last_run = runs[last_index]

        first_offset = max(0, start_pos - first_start)
        last_offset = max(0, end_pos - last_start)

        prefix = (first_run.text or "")[:first_offset]
        suffix = (last_run.text or "")[last_offset:]

        first_run.text = f"{prefix}{replacement}{suffix}"

        for index, _run_start, _run_end in affected[1:]:
            runs[index].text = ""

