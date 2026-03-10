from __future__ import annotations

import os
import re
import time
from io import BytesIO
from typing import Any

import httpx

from models.compliance import ComplianceDocumentChunk
from services.vector_db import VectorDBService


class DocumentIngestionService:
    """Extract text, apply legal-aware chunking, and index into ChromaDB."""

    def __init__(self):
        self.vector_db = VectorDBService()

    def ingest(
        self,
        file_bytes: bytes,
        filename: str,
        collection_name: str,
        document_id: int,
        metadata: dict[str, Any] | None = None,
        db=None,
    ) -> dict[str, Any]:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext == "pdf":
            extraction = self._extract_pdf(file_bytes)
        elif ext == "docx":
            extraction = self._extract_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported format: {ext or 'unknown'}")

        normalized_text = self._normalize_text(extraction["text"])
        if not normalized_text:
            raise ValueError("No extractable text found in uploaded file.")

        base_metadata = dict(metadata or {})
        base_metadata.setdefault("scope", "global")
        base_metadata.setdefault("fund_id", "")
        base_metadata.setdefault("fund_type_filter", "")
        base_metadata.setdefault("investment_id", "")
        base_metadata.setdefault("company_id", "")
        base_metadata.setdefault("source_tier", "law")
        base_metadata.setdefault("document_role", "")

        chunks = self._legal_chunking(
            text=normalized_text,
            document_id=document_id,
            metadata=base_metadata,
        )
        if not chunks:
            raise ValueError("No indexable text chunks were generated.")

        return {
            "chunk_count": len(chunks),
            "ocr_status": extraction["ocr_status"],
            "extraction_quality": extraction["quality"],
            "ingest_status": "indexed",
            "index_status": "indexed",
            **self._index_chunks(
                collection_name=collection_name,
                document_id=document_id,
                chunks=chunks,
                db=db,
            ),
        }

    def ingest_text(
        self,
        *,
        text: str,
        collection_name: str,
        document_id: int,
        metadata: dict[str, Any] | None = None,
        db=None,
    ) -> dict[str, Any]:
        normalized_text = self._normalize_text(text)
        if not normalized_text:
            raise ValueError("No indexable text content found.")

        base_metadata = dict(metadata or {})
        base_metadata.setdefault("scope", "global")
        base_metadata.setdefault("fund_id", "")
        base_metadata.setdefault("fund_type_filter", "")
        base_metadata.setdefault("investment_id", "")
        base_metadata.setdefault("company_id", "")
        base_metadata.setdefault("source_tier", "law")
        base_metadata.setdefault("document_role", "")

        chunks = self._legal_chunking(
            text=normalized_text,
            document_id=document_id,
            metadata=base_metadata,
        )
        if not chunks:
            raise ValueError("No indexable text chunks were generated.")

        return {
            "chunk_count": len(chunks),
            "ocr_status": "not_needed",
            "extraction_quality": self._estimate_text_quality(normalized_text),
            "ingest_status": "indexed",
            "index_status": "indexed",
            **self._index_chunks(
                collection_name=collection_name,
                document_id=document_id,
                chunks=chunks,
                db=db,
            ),
        }

    def _index_chunks(
        self,
        *,
        collection_name: str,
        document_id: int,
        chunks: list[dict[str, Any]],
        db=None,
    ) -> dict[str, Any]:
        self.vector_db.delete_chunks_by_document(collection_name, document_id)
        self.vector_db.add_chunks(collection_name, chunks)
        if db is not None:
            (
                db.query(ComplianceDocumentChunk)
                .filter(ComplianceDocumentChunk.document_id == document_id)
                .delete(synchronize_session=False)
            )
            for item in chunks:
                db.add(
                    ComplianceDocumentChunk(
                        document_id=document_id,
                        chunk_key=str(item["id"]),
                        page_no=item["metadata"].get("page_no"),
                        section_ref=item["metadata"].get("section_ref"),
                        clause_type=item["metadata"].get("clause_type"),
                        chunk_index=int(item["metadata"].get("chunk_index") or 0),
                        token_count=self._estimate_token_count(item["text"]),
                        text=item["text"],
                        metadata_json=item["metadata"],
                    )
                )
        return {}

    @staticmethod
    def _extract_pdf(file_bytes: bytes) -> dict[str, Any]:
        try:
            import pdfplumber
        except Exception as exc:
            raise RuntimeError("pdfplumber is required for PDF ingestion.") from exc
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]

        native_text = DocumentIngestionService._compose_page_text(pages)
        native_quality = DocumentIngestionService._estimate_text_quality(native_text)
        if native_quality >= 0.2:
            return {
                "text": native_text,
                "ocr_status": "not_needed",
                "quality": native_quality,
            }

        azure_endpoint = (os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT") or "").strip().rstrip("/")
        azure_key = (os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY") or "").strip()
        if not azure_endpoint or not azure_key:
            if native_text.strip():
                return {
                    "text": native_text,
                    "ocr_status": "unavailable",
                    "quality": native_quality,
                }
            raise RuntimeError("Scanned PDF detected but Azure Document Intelligence is not configured.")

        try:
            ocr_pages = DocumentIngestionService._extract_pdf_with_azure_ocr(
                file_bytes=file_bytes,
                endpoint=azure_endpoint,
                api_key=azure_key,
            )
            ocr_text = DocumentIngestionService._compose_page_text(ocr_pages)
            ocr_quality = DocumentIngestionService._estimate_text_quality(ocr_text)
            return {
                "text": ocr_text or native_text,
                "ocr_status": "completed",
                "quality": max(native_quality, ocr_quality),
            }
        except Exception:
            if native_text.strip():
                return {
                    "text": native_text,
                    "ocr_status": "failed",
                    "quality": native_quality,
                }
            raise

    @staticmethod
    def _extract_docx(file_bytes: bytes) -> dict[str, Any]:
        try:
            from docx import Document
        except Exception as exc:
            raise RuntimeError("python-docx is required for DOCX ingestion.") from exc
        doc = Document(BytesIO(file_bytes))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)
        return {
            "text": text,
            "ocr_status": "not_needed",
            "quality": DocumentIngestionService._estimate_text_quality(text),
        }

    @staticmethod
    def _normalize_text(text: str) -> str:
        lines = [line.rstrip() for line in text.splitlines()]
        compact = "\n".join(lines)
        compact = re.sub(r"\n{3,}", "\n\n", compact)
        return compact.strip()

    @staticmethod
    def _compose_page_text(pages: list[str]) -> str:
        parts: list[str] = []
        for page_no, page_text in enumerate(pages, start=1):
            normalized = (page_text or "").strip()
            if not normalized:
                continue
            parts.append(f"[[PAGE:{page_no}]]\n{normalized}")
        return "\n\n".join(parts)

    @staticmethod
    def _estimate_text_quality(text: str) -> float:
        normalized = (text or "").strip()
        if not normalized:
            return 0.0
        visible_chars = len(re.sub(r"\s+", "", normalized))
        alpha_num_chars = len(re.findall(r"[0-9A-Za-z가-힣]", normalized))
        if visible_chars <= 0:
            return 0.0
        density = min(1.0, alpha_num_chars / max(visible_chars, 1))
        length_factor = min(1.0, visible_chars / 2000)
        return round((density * 0.6) + (length_factor * 0.4), 4)

    @staticmethod
    def _estimate_token_count(text: str) -> int:
        return max(1, int(len((text or "").strip()) / 4))

    @staticmethod
    def _extract_page_no(text: str) -> int | None:
        matched = re.search(r"\[\[PAGE:(\d+)\]\]", text)
        if not matched:
            return None
        return int(matched.group(1))

    @staticmethod
    def _strip_page_markers(text: str) -> str:
        return re.sub(r"\[\[PAGE:\d+\]\]\s*", "", text).strip()

    @staticmethod
    def _extract_pdf_with_azure_ocr(
        *,
        file_bytes: bytes,
        endpoint: str,
        api_key: str,
    ) -> list[str]:
        api_version = "2024-02-29-preview"
        analyze_url = f"{endpoint}/documentintelligence/documentModels/prebuilt-layout:analyze?api-version={api_version}"
        headers = {
            "Ocp-Apim-Subscription-Key": api_key,
            "Content-Type": "application/pdf",
        }

        with httpx.Client(timeout=60.0) as client:
            response = client.post(analyze_url, headers=headers, content=file_bytes)
            response.raise_for_status()
            operation_url = response.headers.get("Operation-Location")
            if not operation_url:
                raise RuntimeError("Azure OCR response missing Operation-Location header.")

            for _ in range(20):
                poll_response = client.get(operation_url, headers={"Ocp-Apim-Subscription-Key": api_key})
                poll_response.raise_for_status()
                payload = poll_response.json()
                status = str(payload.get("status") or "").lower()
                if status == "succeeded":
                    pages = payload.get("analyzeResult", {}).get("pages", []) or []
                    extracted_pages: list[str] = []
                    for page in pages:
                        lines = page.get("lines", []) or []
                        page_text = "\n".join(str(line.get("content") or "").strip() for line in lines if str(line.get("content") or "").strip())
                        extracted_pages.append(page_text)
                    return extracted_pages
                if status == "failed":
                    raise RuntimeError("Azure OCR analysis failed.")
                time.sleep(1)
            raise RuntimeError("Azure OCR analysis timed out.")

    @staticmethod
    def _detect_clause_type(text: str) -> str | None:
        lowered = (text or "").lower()
        if any(token in lowered for token in ["보고", "통지", "공시", "신고"]):
            return "reporting"
        if any(token in lowered for token in ["투자한도", "한도", "비율"]):
            return "limit"
        if any(token in lowered for token in ["승인", "동의", "결의", "의결"]):
            return "approval"
        if any(token in lowered for token in ["진술", "보장", "보증"]):
            return "representation"
        if any(token in lowered for token in ["청산", "상환", "회수", "분배"]):
            return "exit"
        return None

    def _legal_chunking(
        self,
        text: str,
        document_id: int,
        metadata: dict[str, Any],
    ) -> list[dict[str, Any]]:
        article_pattern = re.compile(r"(?m)^\s*제\s*\d+\s*조(?:\s*의\s*\d+)?")
        if article_pattern.search(text):
            article_chunks = self._chunk_by_articles(text, document_id, metadata)
            if article_chunks:
                return article_chunks
        return self._chunk_by_sliding_window(
            text=text,
            doc_id=document_id,
            metadata=metadata,
            chunk_size=1000,
            overlap=200,
        )

    def _chunk_by_articles(
        self,
        text: str,
        doc_id: int,
        metadata: dict[str, Any],
    ) -> list[dict[str, Any]]:
        heading_pattern = re.compile(
            r"(?m)^\s*(제\s*\d+\s*조(?:\s*의\s*\d+)?(?:\s*[（(][^）)]*[）)])?)\s*"
        )
        matches = list(heading_pattern.finditer(text))
        if not matches:
            return []

        chunks: list[dict[str, Any]] = []
        for index, match in enumerate(matches):
            start = match.start()
            end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            block = text[start:end].strip()
            if len(block) < 100:
                continue

            article = match.group(1).strip()
            page_no = self._extract_page_no(block)
            clean_block = self._strip_page_markers(block)

            # Sub-chunk long articles to keep within optimal embedding size
            if len(clean_block) > 2000:
                sub_chunks = self._chunk_by_sliding_window(
                    text=clean_block,
                    doc_id=doc_id,
                    metadata={
                        **metadata,
                        "article": article,
                        "page_no": page_no,
                        "section_ref": article,
                        "clause_type": self._detect_clause_type(clean_block),
                    },
                    chunk_size=1000,
                    overlap=200,
                )
                for sc_idx, sc in enumerate(sub_chunks):
                    sc["id"] = f"doc{doc_id}_art_{len(chunks)}"
                    sc["metadata"]["chunk_index"] = len(chunks)
                    sc["metadata"]["article"] = article
                    sc["metadata"]["sub_chunk"] = sc_idx
                    chunks.append(sc)
            else:
                chunks.append(
                    {
                        "id": f"doc{doc_id}_art_{len(chunks)}",
                        "text": clean_block,
                        "metadata": {
                            **metadata,
                            "document_id": doc_id,
                            "article": article,
                            "page_no": page_no,
                            "section_ref": article,
                            "clause_type": self._detect_clause_type(clean_block),
                            "chunk_index": len(chunks),
                        },
                    }
                )
        return chunks

    @staticmethod
    def _find_sentence_boundary(text: str, pos: int, search_range: int = 100) -> int:
        """Find the nearest sentence boundary (period, newline) near *pos*."""
        if pos >= len(text):
            return len(text)
        # Search backward from pos within search_range
        search_start = max(0, pos - search_range)
        segment = text[search_start:pos]
        for delim in (". ", ".\n", "\n\n", "\n"):
            idx = segment.rfind(delim)
            if idx != -1:
                return search_start + idx + len(delim)
        return pos

    def _chunk_by_sliding_window(
        self,
        text: str,
        doc_id: int,
        metadata: dict[str, Any],
        chunk_size: int,
        overlap: int,
    ) -> list[dict[str, Any]]:
        if chunk_size <= overlap:
            raise ValueError("chunk_size must be greater than overlap")

        chunks: list[dict[str, Any]] = []
        start = 0
        step = chunk_size - overlap
        while start < len(text):
            raw_end = min(start + chunk_size, len(text))
            # Snap to sentence boundary unless we are at the end
            end = self._find_sentence_boundary(text, raw_end) if raw_end < len(text) else raw_end
            chunk_text = self._strip_page_markers(text[start:end])
            if len(chunk_text) > 100:
                chunks.append(
                    {
                        "id": f"doc{doc_id}_sw_{len(chunks)}",
                        "text": chunk_text,
                        "metadata": {
                            **metadata,
                            "document_id": doc_id,
                            "page_no": self._extract_page_no(text[start:end]),
                            "section_ref": metadata.get("article") or f"chunk-{len(chunks)}",
                            "clause_type": self._detect_clause_type(chunk_text),
                            "chunk_index": len(chunks),
                        },
                    }
                )
            start += step
        return chunks
