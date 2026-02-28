from __future__ import annotations

import asyncio
import csv
import os
import platform
import multiprocessing as mp
import re
import shutil
import subprocess
import threading
from pathlib import Path
from typing import Callable

try:
    import win32com.client  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    win32com = None  # type: ignore
try:
    import pythoncom  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    pythoncom = None  # type: ignore

from docx import Document

ProgressCallback = Callable[[int, int, str], None]

_HWP_LOCK = threading.Lock()
_STAGE_PREFIX_PATTERN = re.compile(r"^\s*([1-5])\.")
_INVALID_PATH_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1F]')


def _is_windows_hwp_available() -> bool:
    if os.getenv("DOCUMENT_GENERATION_DISABLE_HWP", "").strip().lower() in {"1", "true", "yes", "on"}:
        return False
    return platform.system().lower().startswith("win") and win32com is not None


class HwpAutomation:
    """HWP COM 래퍼."""

    def __init__(self, visible: bool = False):
        self.hwp = None
        self.visible = visible
        self._com_initialized = False

    def __enter__(self):
        if not _is_windows_hwp_available():
            raise RuntimeError("HWP COM is only available on Windows with pywin32 installed.")
        if pythoncom is not None:
            pythoncom.CoInitialize()
            self._com_initialized = True
        self.hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
        try:
            self.hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        except Exception:
            pass
        if not self.visible:
            try:
                self.hwp.XHwpWindows.Item(0).Visible = False
            except Exception:
                pass
        return self

    def __exit__(self, *args):
        if self.hwp:
            try:
                self.hwp.Quit()
            except Exception:
                pass
            self.hwp = None
        if self._com_initialized and pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            self._com_initialized = False

    def replace_in_file(self, src_path: str, dst_path: str, replacements: dict) -> tuple[bool, list[str]]:
        abs_src = str(Path(src_path).resolve())
        abs_dst = str(Path(dst_path).resolve())
        Path(abs_dst).parent.mkdir(parents=True, exist_ok=True)

        if not self.hwp.Open(abs_src, "HWP", "forceopen:true"):
            return False, list(replacements.keys())

        not_found: list[str] = []
        for marker, value in replacements.items():
            marker_str = f"{{{{{marker}}}}}"
            found = self._find_replace(marker_str, str(value))
            if not found:
                not_found.append(marker)

        self.hwp.SaveAs(abs_dst, "HWP")
        self.hwp.Clear(1)
        return True, not_found

    def _find_replace(self, find_str: str, replace_str: str) -> bool:
        try:
            act = self.hwp.CreateAction("AllReplace")
            pset = act.CreateSet()
            act.GetDefault(pset)
            pset.SetItem("FindString", find_str)
            pset.SetItem("ReplaceString", replace_str)
            pset.SetItem("IgnoreCase", 0)
            pset.SetItem("WholeWordOnly", 0)
            pset.SetItem("AllWordForms", 0)
            pset.SetItem("SeveralWords", 0)
            pset.SetItem("UseWildCards", 0)
            pset.SetItem("FindRegExp", 0)
            pset.SetItem("ReplaceMode", 1)
            pset.SetItem("FindJaso", 0)
            pset.SetItem("HanjaFromHangul", 0)
            result = act.Execute(pset)
            return result != 0
        except Exception:
            return False


def replace_in_docx(src_path: str, dst_path: str, replacements: dict) -> tuple[bool, list[str]]:
    try:
        Path(dst_path).parent.mkdir(parents=True, exist_ok=True)
        doc = Document(src_path)
        found_keys: set[str] = set()

        def _replace_in_runs(runs):
            full_text = "".join(r.text for r in runs)
            for marker, value in replacements.items():
                marker_str = f"{{{{{marker}}}}}"
                if marker_str in full_text:
                    full_text = full_text.replace(marker_str, str(value))
                    found_keys.add(marker)
            if runs:
                runs[0].text = full_text
                for run in runs[1:]:
                    run.text = ""

        for paragraph in doc.paragraphs:
            _replace_in_runs(paragraph.runs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_runs(paragraph.runs)

        not_found = [key for key in replacements if key not in found_keys]
        doc.save(dst_path)
        return True, not_found
    except Exception as exc:
        return False, [str(exc)]


def _resolve_stage(relative_path: Path) -> int | None:
    if not relative_path.parts:
        return None
    match = _STAGE_PREFIX_PATTERN.match(relative_path.parts[0].strip())
    if not match:
        return None
    return int(match.group(1))


def _should_include(relative_path: Path, stages: set[int] | None) -> bool:
    if stages is None:
        return True
    stage = _resolve_stage(relative_path)
    return stage in stages if stage is not None else False


def _replace_markers_in_path(path: Path, replacements: dict) -> Path:
    def _sanitize_component(value: str) -> str:
        # Windows 경로 금지 문자를 제거해 드라이브 경로 오해석(V: 등)을 방지한다.
        sanitized = _INVALID_PATH_CHARS.sub("_", value).strip()
        sanitized = sanitized.rstrip(" .")
        return sanitized or "untitled"

    replaced_parts: list[str] = []
    for part in path.parts:
        next_part = part
        for marker, value in replacements.items():
            next_part = next_part.replace(f"{{{{{marker}}}}}", str(value))
        replaced_parts.append(_sanitize_component(next_part))
    return Path(*replaced_parts)


def _list_hwp_pids_windows() -> set[int]:
    if not platform.system().lower().startswith("win"):
        return set()
    try:
        result = subprocess.run(
            ["tasklist", "/FI", "IMAGENAME eq Hwp.exe", "/FO", "CSV", "/NH"],
            check=False,
            capture_output=True,
            text=True,
        )
    except Exception:
        return set()
    if not result.stdout:
        return set()

    pids: set[int] = set()
    reader = csv.reader(line for line in result.stdout.splitlines() if line.strip())
    for row in reader:
        if len(row) < 2:
            continue
        image_name = (row[0] or "").strip().lower()
        if image_name != "hwp.exe":
            continue
        pid_text = (row[1] or "").replace(",", "").strip()
        try:
            pids.add(int(pid_text))
        except ValueError:
            continue
    return pids


def _run_hwp_worker(
    src_path: str,
    dst_path: str,
    replacements: dict[str, str],
    result_queue,
) -> None:
    try:
        with HwpAutomation(visible=False) as hwp_auto:
            ok, not_found = hwp_auto.replace_in_file(src_path, dst_path, replacements)
        result_queue.put({"ok": bool(ok), "not_found": list(not_found), "error": None})
    except Exception as exc:
        result_queue.put({"ok": False, "not_found": [], "error": str(exc)})


def _replace_hwp_with_timeout(
    src_path: str,
    dst_path: str,
    replacements: dict[str, str],
    timeout_sec: int,
) -> tuple[bool, list[str], str | None, bool]:
    hwp_before = _list_hwp_pids_windows()
    ctx = mp.get_context("spawn")
    result_queue = ctx.Queue()
    proc = ctx.Process(target=_run_hwp_worker, args=(src_path, dst_path, replacements, result_queue))
    proc.start()
    proc.join(timeout=timeout_sec)

    if proc.is_alive():
        # 타임아웃 시 워커 프로세스 트리(HWP 포함)를 강제 종료해 잔류 프로세스를 방지한다.
        try:
            subprocess.run(
                ["taskkill", "/PID", str(proc.pid), "/T", "/F"],
                check=False,
                capture_output=True,
                text=True,
            )
        except Exception:
            proc.terminate()
        proc.join(timeout=5)
        hwp_after = _list_hwp_pids_windows()
        spawned_hwp = hwp_after - hwp_before
        for pid in spawned_hwp:
            try:
                subprocess.run(
                    ["taskkill", "/PID", str(pid), "/F"],
                    check=False,
                    capture_output=True,
                    text=True,
                )
            except Exception:
                pass
        return False, [], f"HWP 처리 타임아웃({timeout_sec}초)", True

    if not result_queue.empty():
        payload = result_queue.get()
        return bool(payload.get("ok")), list(payload.get("not_found") or []), payload.get("error"), False

    if proc.exitcode != 0:
        return False, [], f"HWP 워커 비정상 종료(exitcode={proc.exitcode})", False
    return False, [], "HWP 처리 결과를 수신하지 못했습니다.", False


def _safe_progress(callback: ProgressCallback | None, current: int, total: int, message: str) -> None:
    if callback is None:
        return
    try:
        callback(current, total, message)
    except Exception:
        pass


def _generate_documents_sync(
    variables: dict,
    template_dir: str,
    output_dir: str,
    stages: list[int] | None = None,
    progress_callback: ProgressCallback | None = None,
) -> dict:
    template_path = Path(template_dir).resolve()
    if not template_path.exists():
        raise FileNotFoundError(f"템플릿 폴더를 찾을 수 없습니다: {template_path}")

    selected_stages = set(stages) if stages else None
    if selected_stages:
        for stage in selected_stages:
            if stage not in {1, 2, 3, 4, 5}:
                raise ValueError("stages must contain values between 1 and 5.")

    raw_fund_name = str(variables.get("조합명", "새조합") or "").strip() or "새조합"
    fund_name = _INVALID_PATH_CHARS.sub("_", raw_fund_name).strip().rstrip(" .") or "새조합"
    output_base = Path(output_dir).resolve() / fund_name
    output_base.mkdir(parents=True, exist_ok=True)

    selected_files: list[Path] = []
    for item in template_path.rglob("*"):
        if not item.is_file():
            continue
        if item.name == ".gitkeep":
            continue
        relative_path = item.relative_to(template_path)
        if _should_include(relative_path, selected_stages):
            selected_files.append(item)
    selected_files.sort(key=lambda row: str(row.relative_to(template_path)).lower())

    hwp_files = [row for row in selected_files if row.suffix.lower() == ".hwp"]
    docx_files = [row for row in selected_files if row.suffix.lower() == ".docx"]
    copy_files = [row for row in selected_files if row.suffix.lower() not in {".hwp", ".docx"}]

    total = len(selected_files)
    current = 0
    results = {"success": [], "failed": [], "warnings": [], "total_files": total, "output_path": str(output_base)}

    if hwp_files:
        if _is_windows_hwp_available():
            timeout_raw = os.getenv("DOCUMENT_GENERATION_HWP_FILE_TIMEOUT_SEC", "15").strip()
            try:
                hwp_timeout_sec = max(10, int(timeout_raw))
            except ValueError:
                hwp_timeout_sec = 15

            with _HWP_LOCK:
                for file_path in hwp_files:
                    current += 1
                    relative = file_path.relative_to(template_path)
                    replaced_relative = _replace_markers_in_path(relative, variables)
                    destination = output_base / replaced_relative
                    _safe_progress(progress_callback, current, total, f"처리 중: {relative.name}")
                    ok, not_found, error, timed_out = _replace_hwp_with_timeout(
                        str(file_path),
                        str(destination),
                        {str(k): str(v) for k, v in variables.items()},
                        hwp_timeout_sec,
                    )
                    if ok:
                        results["success"].append(str(replaced_relative))
                        if not_found:
                            results["warnings"].append(f"{relative.name}: 미치환 마커 {not_found}")
                        continue

                    # HWP COM 실패/멈춤 시 생성 중단 대신 원본 복사로 진행을 유지한다.
                    try:
                        destination.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copy2(file_path, destination)
                        results["success"].append(f"{replaced_relative} (복사)")
                        reason = error or "알 수 없는 HWP 처리 오류"
                        if timed_out:
                            reason = f"{reason} - HWP COM 응답 지연"
                        results["warnings"].append(f"{relative.name}: HWP 치환 실패로 원본 복사 ({reason})")
                    except Exception:
                        results["failed"].append(str(replaced_relative))
        else:
            for file_path in hwp_files:
                current += 1
                relative = file_path.relative_to(template_path)
                replaced_relative = _replace_markers_in_path(relative, variables)
                destination = output_base / replaced_relative
                destination.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(file_path, destination)
                _safe_progress(progress_callback, current, total, f"HWP 복사: {relative.name}")
                results["success"].append(f"{replaced_relative} (복사)")
                results["warnings"].append(f"{relative.name}: Windows/HWP COM 미지원으로 원본 복사")

    for file_path in docx_files:
        current += 1
        relative = file_path.relative_to(template_path)
        replaced_relative = _replace_markers_in_path(relative, variables)
        destination = output_base / replaced_relative
        _safe_progress(progress_callback, current, total, f"처리 중: {relative.name}")
        ok, not_found = replace_in_docx(str(file_path), str(destination), variables)
        if ok:
            results["success"].append(str(replaced_relative))
            if not_found:
                results["warnings"].append(f"{relative.name}: 미치환 마커 {not_found}")
        else:
            results["failed"].append(str(replaced_relative))

    for file_path in copy_files:
        current += 1
        relative = file_path.relative_to(template_path)
        replaced_relative = _replace_markers_in_path(relative, variables)
        destination = output_base / replaced_relative
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(file_path, destination)
        _safe_progress(progress_callback, current, total, f"복사: {relative.name}")
        results["success"].append(f"{replaced_relative} (복사)")

    _safe_progress(progress_callback, total, total, "완료!")
    return results


async def generate_documents(
    variables: dict,
    template_dir: str,
    output_dir: str,
    stages: list[int] | None = None,
    progress_callback: ProgressCallback | None = None,
) -> dict:
    return await asyncio.to_thread(
        _generate_documents_sync,
        variables,
        template_dir,
        output_dir,
        stages,
        progress_callback,
    )
