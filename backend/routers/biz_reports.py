from __future__ import annotations

import base64
from datetime import date, datetime, timedelta
from decimal import Decimal
from io import BytesIO
from typing import Literal

from docx import Document
from fastapi import APIRouter, Body, Depends, HTTPException
from openpyxl import Workbook
from pydantic import BaseModel, Field
from sqlalchemy import func
from sqlalchemy.orm import Session

from database import get_db
from models.biz_report import BizReport, BizReportAnomaly, BizReportRequest, BizReportTemplate
from models.fee import ManagementFee, PerformanceFeeSimulation
from models.fund import Fund
from models.investment import Investment, PortfolioCompany
from models.phase3 import Distribution
from models.transaction import Transaction
from models.valuation import Valuation
from schemas.biz_report import (
    BizReportAnomalyResponse,
    BizReportCommentDiffResponse,
    BizReportCreate,
    BizReportGenerationResponse,
    BizReportMatrixCell,
    BizReportMatrixResponse,
    BizReportMatrixRow,
    BizReportRequestResponse,
    BizReportRequestUpdate,
    BizReportResponse,
    BizReportTemplateCreate,
    BizReportTemplateResponse,
    BizReportUpdate,
)
from services.generated_document_service import generate_and_store_document
from services.biz_report_anomaly import detect_biz_report_anomalies
from services.biz_report_valuation_sync import suggest_valuation_updates

router = APIRouter(tags=["biz-reports"])

NUMERIC_FIELDS = [
    "total_commitment",
    "total_paid_in",
    "total_invested",
    "total_distributed",
    "fund_nav",
    "irr",
    "tvpi",
    "dpi",
]
DATE_FIELDS = ["submission_date", "created_at"]
DOC_STATUS_MAP = {
    "financial_statement": "doc_financial_statement",
    "biz_registration": "doc_biz_registration",
    "shareholder_list": "doc_shareholder_list",
    "corp_registry": "doc_corp_registry",
    "insurance_cert": "doc_insurance_cert",
    "credit_report": "doc_credit_report",
    "other_changes": "doc_other_changes",
}
DOC_STATUS_ALLOWED = {"not_requested", "requested", "received", "verified"}
DOC_RECEIVED_STATES = {"received", "verified"}


class BizReportRequestDocStatusPatch(BaseModel):
    doc_type: Literal[
        "financial_statement",
        "biz_registration",
        "shareholder_list",
        "corp_registry",
        "insurance_cert",
        "credit_report",
        "other_changes",
    ]
    status: Literal["not_requested", "requested", "received", "verified"]


class BizReportSendDocRequestsBody(BaseModel):
    quarter: int | None = Field(default=None, ge=1, le=4)
    deadline_days: int = Field(default=14, ge=1, le=90)


def _to_primitive(value):
    if isinstance(value, Decimal):
        return float(value)
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    return value


def _serialize_report(report: BizReport, fund_name: str | None = None) -> dict:
    data = {
        column.name: _to_primitive(getattr(report, column.name))
        for column in report.__table__.columns
    }
    data["fund_name"] = fund_name

    for field in NUMERIC_FIELDS:
        value = data.get(field)
        if value is not None:
            data[field] = float(value)

    for field in DATE_FIELDS:
        if field in data and data[field] is not None:
            data[field] = str(data[field])

    return data


def _serialize_request(db: Session, row: BizReportRequest) -> BizReportRequestResponse:
    investment = db.get(Investment, row.investment_id)
    company_name = None
    if investment is not None:
        company = db.get(PortfolioCompany, investment.company_id)
        company_name = company.name if company else f"Investment #{row.investment_id}"
    return BizReportRequestResponse(
        id=row.id,
        biz_report_id=row.biz_report_id,
        investment_id=row.investment_id,
        investment_name=company_name,
        request_date=row.request_date,
        deadline=row.deadline,
        status=row.status,
        revenue=float(row.revenue) if row.revenue is not None else None,
        operating_income=float(row.operating_income) if row.operating_income is not None else None,
        net_income=float(row.net_income) if row.net_income is not None else None,
        total_assets=float(row.total_assets) if row.total_assets is not None else None,
        total_equity=float(row.total_equity) if row.total_equity is not None else None,
        cash=float(row.cash) if row.cash is not None else None,
        employees=row.employees,
        prev_revenue=float(row.prev_revenue) if row.prev_revenue is not None else None,
        prev_operating_income=float(row.prev_operating_income) if row.prev_operating_income is not None else None,
        prev_net_income=float(row.prev_net_income) if row.prev_net_income is not None else None,
        doc_financial_statement=row.doc_financial_statement or "not_requested",
        doc_biz_registration=row.doc_biz_registration or "not_requested",
        doc_shareholder_list=row.doc_shareholder_list or "not_requested",
        doc_corp_registry=row.doc_corp_registry or "not_requested",
        doc_insurance_cert=row.doc_insurance_cert or "not_requested",
        doc_credit_report=row.doc_credit_report or "not_requested",
        doc_other_changes=row.doc_other_changes or "not_requested",
        request_sent_date=row.request_sent_date,
        request_deadline=row.request_deadline,
        all_docs_received_date=row.all_docs_received_date,
        comment=row.comment,
        reviewer_comment=row.reviewer_comment,
        risk_flag=row.risk_flag,
        created_at=row.created_at,
        updated_at=row.updated_at,
    )


def _serialize_anomaly(row: BizReportAnomaly) -> BizReportAnomalyResponse:
    return BizReportAnomalyResponse(
        id=row.id,
        request_id=row.request_id,
        anomaly_type=row.anomaly_type,
        severity=row.severity,
        detail=row.detail,
        acknowledged=bool(row.acknowledged),
        created_at=row.created_at,
    )


def _doc_status_payload(row: BizReportRequest) -> dict[str, str]:
    return {
        "financial_statement": row.doc_financial_statement or "not_requested",
        "biz_registration": row.doc_biz_registration or "not_requested",
        "shareholder_list": row.doc_shareholder_list or "not_requested",
        "corp_registry": row.doc_corp_registry or "not_requested",
        "insurance_cert": row.doc_insurance_cert or "not_requested",
        "credit_report": row.doc_credit_report or "not_requested",
        "other_changes": row.doc_other_changes or "not_requested",
    }


def _received_doc_count(row: BizReportRequest) -> int:
    payload = _doc_status_payload(row)
    return sum(1 for value in payload.values() if value in DOC_RECEIVED_STATES)


def _is_all_docs_received(row: BizReportRequest) -> bool:
    return _received_doc_count(row) >= len(DOC_STATUS_MAP)


def _report_quarter(report: BizReport) -> int:
    base = report.created_at.date() if report.created_at else date.today()
    return ((base.month - 1) // 3) + 1


def _latest_nav(db: Session, fund_id: int) -> float:
    rows = (
        db.query(Valuation)
        .filter(Valuation.fund_id == fund_id)
        .order_by(Valuation.as_of_date.desc(), Valuation.id.desc())
        .all()
    )
    latest_by_investment: dict[int, Valuation] = {}
    for row in rows:
        if row.investment_id not in latest_by_investment:
            latest_by_investment[row.investment_id] = row
    return float(sum(float(row.total_fair_value or row.value or 0) for row in latest_by_investment.values()))


def _transaction_summary(db: Session, fund_id: int) -> tuple[float, float]:
    invested = float(
        db.query(func.coalesce(func.sum(Transaction.amount), 0))
        .filter(Transaction.fund_id == fund_id, Transaction.type == "investment")
        .scalar()
        or 0
    )
    recovered = float(
        db.query(func.coalesce(func.sum(Transaction.amount), 0))
        .filter(Transaction.fund_id == fund_id, Transaction.type == "exit")
        .scalar()
        or 0
    )
    return invested, recovered


@router.get("/api/biz-reports", response_model=list[BizReportResponse])
def list_biz_reports(
    fund_id: int | None = None,
    year: int | None = None,
    status: str | None = None,
    db: Session = Depends(get_db),
):
    query = db.query(BizReport)
    if fund_id:
        query = query.filter(BizReport.fund_id == fund_id)
    if year:
        query = query.filter(BizReport.report_year == year)
    if status:
        query = query.filter(BizReport.status == status)

    reports = query.order_by(BizReport.report_year.desc(), BizReport.created_at.desc(), BizReport.id.desc()).all()
    result = []
    for report in reports:
        fund = db.get(Fund, report.fund_id)
        result.append(_serialize_report(report, fund.name if fund else None))
    return result


@router.get("/api/biz-reports/{report_id}", response_model=BizReportResponse)
def get_biz_report(report_id: int, db: Session = Depends(get_db)):
    report = db.get(BizReport, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")
    fund = db.get(Fund, report.fund_id)
    return _serialize_report(report, fund.name if fund else None)


@router.get("/api/biz-reports/{report_id}/valuation-suggestions")
async def get_valuation_suggestions(report_id: int, db: Session = Depends(get_db)):
    try:
        return await suggest_valuation_updates(db, report_id)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@router.post("/api/biz-reports", response_model=BizReportResponse, status_code=201)
def create_biz_report(data: BizReportCreate, db: Session = Depends(get_db)):
    fund = db.get(Fund, data.fund_id)
    if not fund:
        raise HTTPException(status_code=404, detail="조합을 찾을 수 없습니다")

    report = BizReport(**data.model_dump())
    db.add(report)
    try:
        db.commit()
    except Exception:
        db.rollback()
        raise
    db.refresh(report)
    return _serialize_report(report, fund.name)


@router.put("/api/biz-reports/{report_id}", response_model=BizReportResponse)
def update_biz_report(report_id: int, data: BizReportUpdate, db: Session = Depends(get_db)):
    report = db.get(BizReport, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")

    payload = data.model_dump(exclude_unset=True)
    next_fund_id = payload.get("fund_id", report.fund_id)
    fund = db.get(Fund, next_fund_id)
    if not fund:
        raise HTTPException(status_code=404, detail="조합을 찾을 수 없습니다")

    for key, value in payload.items():
        setattr(report, key, value)

    try:
        db.commit()
    except Exception:
        db.rollback()
        raise
    db.refresh(report)
    return _serialize_report(report, fund.name)


@router.delete("/api/biz-reports/{report_id}")
def delete_biz_report(report_id: int, db: Session = Depends(get_db)):
    report = db.get(BizReport, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")
    db.delete(report)
    try:
        db.commit()
    except Exception:
        db.rollback()
        raise
    return {"ok": True}


@router.get("/api/biz-reports/matrix", response_model=BizReportMatrixResponse)
def get_biz_report_matrix(year: int | None = None, db: Session = Depends(get_db)):
    target_year = year or date.today().year
    funds = db.query(Fund).order_by(Fund.name.asc(), Fund.id.asc()).all()
    report_rows = db.query(BizReport).filter(BizReport.report_year == target_year).all()

    by_fund: dict[int, dict[str, BizReport]] = {}
    for report in report_rows:
        created_at = report.created_at or datetime(target_year, 1, 1)
        quarter = f"{((created_at.month - 1) // 3) + 1}Q"
        current = by_fund.setdefault(report.fund_id, {}).get(quarter)
        if current is None or (current.created_at or datetime.min) < created_at:
            by_fund.setdefault(report.fund_id, {})[quarter] = report

    rows: list[BizReportMatrixRow] = []
    for fund in funds:
        cells: list[BizReportMatrixCell] = []
        for quarter in ["1Q", "2Q", "3Q", "4Q"]:
            report = by_fund.get(fund.id, {}).get(quarter)
            cells.append(
                BizReportMatrixCell(
                    quarter=quarter,
                    status=report.status if report else "예정",
                    report_id=report.id if report else None,
                )
            )
        rows.append(BizReportMatrixRow(fund_id=fund.id, fund_name=fund.name, cells=cells))
    return BizReportMatrixResponse(rows=rows)


@router.get("/api/biz-reports/{report_id}/requests", response_model=list[BizReportRequestResponse])
def list_biz_report_requests(report_id: int, db: Session = Depends(get_db)):
    report = db.get(BizReport, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")
    rows = (
        db.query(BizReportRequest)
        .filter(BizReportRequest.biz_report_id == report_id)
        .order_by(BizReportRequest.id.asc())
        .all()
    )
    return [_serialize_request(db, row) for row in rows]


@router.get("/api/biz-reports/doc-collection-matrix")
def get_doc_collection_matrix(
    report_id: int | None = None,
    fund_id: int | None = None,
    year: int | None = None,
    quarter: int | None = None,
    db: Session = Depends(get_db),
):
    report: BizReport | None = None
    if report_id is not None:
        report = db.get(BizReport, report_id)
        if not report:
            raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")
    else:
        query = db.query(BizReport)
        if fund_id is not None:
            query = query.filter(BizReport.fund_id == fund_id)
        if year is not None:
            query = query.filter(BizReport.report_year == year)
        candidates = query.order_by(BizReport.created_at.desc(), BizReport.id.desc()).all()
        if quarter is not None:
            candidates = [row for row in candidates if _report_quarter(row) == quarter]
        report = candidates[0] if candidates else None

    if not report:
        raise HTTPException(status_code=404, detail="조회 가능한 영업보고가 없습니다")

    requests = (
        db.query(BizReportRequest)
        .filter(BizReportRequest.biz_report_id == report.id)
        .order_by(BizReportRequest.id.asc())
        .all()
    )
    fund = db.get(Fund, report.fund_id)

    companies: list[dict] = []
    completed_companies = 0
    for req in requests:
        investment = db.get(Investment, req.investment_id)
        company = db.get(PortfolioCompany, investment.company_id) if investment else None
        docs = _doc_status_payload(req)
        received_count = _received_doc_count(req)
        completed = received_count >= len(DOC_STATUS_MAP)
        if completed:
            completed_companies += 1

        if completed:
            status_label = "완료"
        elif received_count > 0:
            status_label = f"{received_count}/{len(DOC_STATUS_MAP)}"
        elif any(value == "requested" for value in docs.values()):
            status_label = "요청중"
        else:
            status_label = "미요청"

        companies.append(
            {
                "company_name": company.name if company else f"Investment #{req.investment_id}",
                "request_id": req.id,
                "docs": docs,
                "received_count": received_count,
                "status": status_label,
            }
        )

    total_companies = len(companies)
    completion_pct = round((completed_companies / total_companies * 100), 1) if total_companies else 0.0
    current_quarter = f"{report.report_year}-Q{_report_quarter(report)}"
    return {
        "report_id": report.id,
        "fund_id": report.fund_id,
        "fund_name": fund.name if fund else f"Fund #{report.fund_id}",
        "quarter": current_quarter,
        "total_companies": total_companies,
        "completed_companies": completed_companies,
        "completion_pct": completion_pct,
        "companies": companies,
    }


@router.patch("/api/biz-report-requests/{request_id}/doc-status", response_model=BizReportRequestResponse)
def patch_biz_report_request_doc_status(
    request_id: int,
    data: BizReportRequestDocStatusPatch,
    db: Session = Depends(get_db),
):
    row = db.get(BizReportRequest, request_id)
    if not row:
        raise HTTPException(status_code=404, detail="요청 건을 찾을 수 없습니다")

    if data.status not in DOC_STATUS_ALLOWED:
        raise HTTPException(status_code=400, detail="지원하지 않는 서류 상태입니다")
    column_name = DOC_STATUS_MAP.get(data.doc_type)
    if not column_name:
        raise HTTPException(status_code=400, detail="지원하지 않는 서류 유형입니다")

    setattr(row, column_name, data.status)
    if data.status != "not_requested" and row.request_sent_date is None:
        row.request_sent_date = date.today()
    if row.request_deadline is None and row.request_sent_date is not None:
        row.request_deadline = row.request_sent_date + timedelta(days=14)

    if _is_all_docs_received(row):
        if row.all_docs_received_date is None:
            row.all_docs_received_date = date.today()
        row.status = "완료"
    else:
        row.all_docs_received_date = None
        statuses = list(_doc_status_payload(row).values())
        if any(value in DOC_RECEIVED_STATES for value in statuses):
            row.status = "검토중"
        elif any(value == "requested" for value in statuses):
            row.status = "요청"

    try:
        db.commit()
    except Exception:
        db.rollback()
        raise
    db.refresh(row)
    return _serialize_request(db, row)


@router.post("/api/biz-reports/{report_id}/send-doc-requests")
def send_doc_requests(
    report_id: int,
    body: BizReportSendDocRequestsBody | None = Body(default=None),
    db: Session = Depends(get_db),
):
    report = db.get(BizReport, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")

    payload = body or BizReportSendDocRequestsBody()
    target_quarter = payload.quarter or _report_quarter(report)
    today = date.today()
    deadline = today + timedelta(days=payload.deadline_days)

    rows = (
        db.query(BizReportRequest)
        .filter(BizReportRequest.biz_report_id == report_id)
        .order_by(BizReportRequest.id.asc())
        .all()
    )
    if not rows:
        raise HTTPException(status_code=400, detail="요청 데이터가 없습니다")

    pending_rows: list[BizReportRequest] = []
    for row in rows:
        if _is_all_docs_received(row):
            continue
        for column_name in DOC_STATUS_MAP.values():
            current = getattr(row, column_name, None)
            if current in (None, "", "not_requested"):
                setattr(row, column_name, "requested")
        row.request_sent_date = today
        row.request_deadline = deadline
        row.all_docs_received_date = None
        row.status = "요청"
        pending_rows.append(row)

    try:
        db.commit()
    except Exception:
        db.rollback()
        raise

    generated_documents: list[dict] = []
    failures: list[dict] = []
    for row in pending_rows:
        try:
            generated = generate_and_store_document(
                "doc_request_letter",
                {
                    "fund_id": report.fund_id,
                    "investment_id": row.investment_id,
                    "year": report.report_year,
                    "quarter": target_quarter,
                },
                db,
            )
            generated_documents.append(
                {
                    "request_id": row.id,
                    "investment_id": row.investment_id,
                    **generated,
                }
            )
        except Exception as exc:
            failures.append(
                {
                    "request_id": row.id,
                    "investment_id": row.investment_id,
                    "reason": str(exc),
                }
            )

    return {
        "report_id": report_id,
        "updated_requests": len(pending_rows),
        "generated_documents": generated_documents,
        "failed_documents": failures,
    }


@router.post("/api/biz-reports/{report_id}/requests/generate", response_model=list[BizReportRequestResponse])
def generate_biz_report_requests(report_id: int, db: Session = Depends(get_db)):
    report = db.get(BizReport, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")

    investments = (
        db.query(Investment)
        .filter(Investment.fund_id == report.fund_id)
        .order_by(Investment.id.asc())
        .all()
    )
    if not investments:
        raise HTTPException(status_code=400, detail="해당 조합에 투자건이 없습니다")

    existing_by_investment = {
        row.investment_id: row
        for row in db.query(BizReportRequest).filter(BizReportRequest.biz_report_id == report_id).all()
    }
    generated: list[BizReportRequest] = []
    for investment in investments:
        row = existing_by_investment.get(investment.id)
        if row is None:
            row = BizReportRequest(
                biz_report_id=report_id,
                investment_id=investment.id,
                request_date=date.today(),
                deadline=date.today() + timedelta(days=14),
                status="요청",
                request_sent_date=date.today(),
                request_deadline=date.today() + timedelta(days=14),
            )
            db.add(row)
            generated.append(row)
    try:
        db.commit()
    except Exception:
        db.rollback()
        raise

    rows = (
        db.query(BizReportRequest)
        .filter(BizReportRequest.biz_report_id == report_id)
        .order_by(BizReportRequest.id.asc())
        .all()
    )
    return [_serialize_request(db, row) for row in rows]


@router.patch("/api/biz-report-requests/{request_id}", response_model=BizReportRequestResponse)
def update_biz_report_request(
    request_id: int,
    data: BizReportRequestUpdate,
    db: Session = Depends(get_db),
):
    row = db.get(BizReportRequest, request_id)
    if not row:
        raise HTTPException(status_code=404, detail="요청 건을 찾을 수 없습니다")

    payload = data.model_dump(exclude_unset=True)
    for key, value in payload.items():
        setattr(row, key, value)

    try:
        db.commit()
    except Exception:
        db.rollback()
        raise
    db.refresh(row)
    return _serialize_request(db, row)


@router.post("/api/biz-report-requests/{request_id}/detect-anomalies", response_model=list[BizReportAnomalyResponse])
def detect_anomalies(request_id: int, db: Session = Depends(get_db)):
    row = db.get(BizReportRequest, request_id)
    if not row:
        raise HTTPException(status_code=404, detail="요청 건을 찾을 수 없습니다")

    previous = (
        db.query(BizReportRequest)
        .join(BizReport, BizReport.id == BizReportRequest.biz_report_id)
        .filter(
            BizReportRequest.investment_id == row.investment_id,
            BizReportRequest.id != row.id,
            BizReport.report_year <= db.get(BizReport, row.biz_report_id).report_year,
        )
        .order_by(BizReport.report_year.desc(), BizReportRequest.id.desc())
        .first()
    )

    detected = detect_biz_report_anomalies(row, previous)

    existing = db.query(BizReportAnomaly).filter(BizReportAnomaly.request_id == request_id).all()
    for old in existing:
        db.delete(old)

    created: list[BizReportAnomaly] = []
    for item in detected:
        anomaly = BizReportAnomaly(
            request_id=request_id,
            anomaly_type=item["anomaly_type"],
            severity=item["severity"],
            detail=item["detail"],
            acknowledged=False,
        )
        db.add(anomaly)
        created.append(anomaly)

    if any(item["severity"] == "위험" for item in detected):
        row.risk_flag = "위험"
    elif any(item["severity"] == "주의" for item in detected):
        row.risk_flag = "주의"
    else:
        row.risk_flag = "정상"

    try:
        db.commit()
    except Exception:
        db.rollback()
        raise

    for item in created:
        db.refresh(item)
    return [_serialize_anomaly(item) for item in created]


@router.get("/api/biz-report-requests/{request_id}/anomalies", response_model=list[BizReportAnomalyResponse])
def list_anomalies(request_id: int, db: Session = Depends(get_db)):
    if not db.get(BizReportRequest, request_id):
        raise HTTPException(status_code=404, detail="요청 건을 찾을 수 없습니다")
    rows = (
        db.query(BizReportAnomaly)
        .filter(BizReportAnomaly.request_id == request_id)
        .order_by(BizReportAnomaly.id.asc())
        .all()
    )
    return [_serialize_anomaly(row) for row in rows]


@router.get("/api/biz-report-requests/{request_id}/comment-diff", response_model=BizReportCommentDiffResponse)
def get_comment_diff(request_id: int, db: Session = Depends(get_db)):
    row = db.get(BizReportRequest, request_id)
    if not row:
        raise HTTPException(status_code=404, detail="요청 건을 찾을 수 없습니다")

    current = (row.comment or "").strip() or None
    previous = (
        db.query(BizReportRequest)
        .join(BizReport, BizReport.id == BizReportRequest.biz_report_id)
        .filter(
            BizReportRequest.investment_id == row.investment_id,
            BizReportRequest.id != row.id,
            BizReport.report_year <= db.get(BizReport, row.biz_report_id).report_year,
        )
        .order_by(BizReport.report_year.desc(), BizReportRequest.id.desc())
        .first()
    )
    prev_comment = (previous.comment or "").strip() if previous and previous.comment else None
    return BizReportCommentDiffResponse(
        current_comment=current,
        previous_comment=prev_comment,
        changed=(current or "") != (prev_comment or ""),
    )


@router.get("/api/biz-report-templates", response_model=list[BizReportTemplateResponse])
def list_biz_report_templates(db: Session = Depends(get_db)):
    rows = db.query(BizReportTemplate).order_by(BizReportTemplate.id.desc()).all()
    return rows


@router.post("/api/biz-report-templates", response_model=BizReportTemplateResponse, status_code=201)
def create_biz_report_template(data: BizReportTemplateCreate, db: Session = Depends(get_db)):
    row = BizReportTemplate(**data.model_dump())
    db.add(row)
    try:
        db.commit()
    except Exception:
        db.rollback()
        raise
    db.refresh(row)
    return row


def _build_generation_metrics(db: Session, report: BizReport) -> dict[str, float]:
    nav = _latest_nav(db, report.fund_id)
    invested, recovered = _transaction_summary(db, report.fund_id)
    distribution_total = float(
        db.query(func.coalesce(func.sum(Distribution.principal_total + Distribution.profit_total), 0))
        .filter(Distribution.fund_id == report.fund_id)
        .scalar()
        or 0
    )
    mgmt_fee_total = float(
        db.query(func.coalesce(func.sum(ManagementFee.fee_amount), 0))
        .filter(ManagementFee.fund_id == report.fund_id)
        .scalar()
        or 0
    )
    perf_latest = (
        db.query(PerformanceFeeSimulation)
        .filter(PerformanceFeeSimulation.fund_id == report.fund_id)
        .order_by(PerformanceFeeSimulation.simulation_date.desc(), PerformanceFeeSimulation.id.desc())
        .first()
    )
    perf_fee = float(perf_latest.carry_amount or 0) if perf_latest else 0.0
    return {
        "nav": nav,
        "invested": invested,
        "recovered": recovered,
        "distribution_total": distribution_total,
        "management_fee_total": mgmt_fee_total,
        "performance_fee": perf_fee,
    }


@router.post("/api/biz-reports/{report_id}/generate-excel", response_model=BizReportGenerationResponse)
def generate_excel(report_id: int, db: Session = Depends(get_db)):
    report = db.get(BizReport, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")
    fund = db.get(Fund, report.fund_id)

    wb = Workbook()
    ws = wb.active
    ws.title = "BizReport"
    ws.append(["항목", "값"])

    metrics = _build_generation_metrics(db, report)
    ws.append(["조합", fund.name if fund else f"Fund #{report.fund_id}"])
    ws.append(["보고연도", report.report_year])
    ws.append(["보고상태", report.status])
    ws.append(["NAV", metrics["nav"]])
    ws.append(["투자원장 누적 투자", metrics["invested"]])
    ws.append(["투자원장 누적 회수", metrics["recovered"]])
    ws.append(["LP 배분 누적", metrics["distribution_total"]])
    ws.append(["관리보수 누적", metrics["management_fee_total"]])
    ws.append(["성과보수(최근 시뮬)", metrics["performance_fee"]])

    ws.append([])
    ws.append(["투자사", "요청상태", "매출", "영업이익", "순이익", "리스크"])
    requests = db.query(BizReportRequest).filter(BizReportRequest.biz_report_id == report_id).all()
    for req in requests:
        inv = db.get(Investment, req.investment_id)
        company = db.get(PortfolioCompany, inv.company_id) if inv else None
        ws.append(
            [
                company.name if company else f"Investment #{req.investment_id}",
                req.status,
                _to_primitive(req.revenue),
                _to_primitive(req.operating_income),
                _to_primitive(req.net_income),
                req.risk_flag,
            ]
        )

    output = BytesIO()
    wb.save(output)
    data = base64.b64encode(output.getvalue()).decode("utf-8")
    return BizReportGenerationResponse(
        filename=f"biz_report_{report_id}.xlsx",
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        base64_data=data,
    )


@router.post("/api/biz-reports/{report_id}/generate-docx", response_model=BizReportGenerationResponse)
def generate_docx(report_id: int, db: Session = Depends(get_db)):
    report = db.get(BizReport, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="영업보고를 찾을 수 없습니다")
    fund = db.get(Fund, report.fund_id)
    metrics = _build_generation_metrics(db, report)

    doc = Document()
    doc.add_heading(f"{fund.name if fund else f'Fund #{report.fund_id}'} 영업보고", level=1)
    doc.add_paragraph(f"보고연도: {report.report_year}")
    doc.add_paragraph(f"상태: {report.status}")
    doc.add_paragraph(f"NAV: {metrics['nav']:,.0f}")
    doc.add_paragraph(f"투자 누적: {metrics['invested']:,.0f} / 회수 누적: {metrics['recovered']:,.0f}")
    doc.add_paragraph(f"관리보수 누적: {metrics['management_fee_total']:,.0f}")
    doc.add_paragraph(f"성과보수(최근): {metrics['performance_fee']:,.0f}")

    requests = db.query(BizReportRequest).filter(BizReportRequest.biz_report_id == report_id).all()
    doc.add_heading("투자사 수집 현황", level=2)
    for req in requests:
        inv = db.get(Investment, req.investment_id)
        company = db.get(PortfolioCompany, inv.company_id) if inv else None
        doc.add_paragraph(
            f"- {company.name if company else req.investment_id}: {req.status} / 매출 {float(req.revenue or 0):,.0f} / 코멘트 {req.comment or '-'}"
        )

    output = BytesIO()
    doc.save(output)
    data = base64.b64encode(output.getvalue()).decode("utf-8")
    return BizReportGenerationResponse(
        filename=f"biz_report_{report_id}.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        base64_data=data,
    )
