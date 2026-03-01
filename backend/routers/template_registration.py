from __future__ import annotations

import json
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from database import get_db
from models.document_template import DocumentTemplate
from models.template_variable import TemplateVariable
from schemas.template_registration import (
    AnalyzeTemplateResponse,
    InputCacheResponse,
    MarkerCandidate,
    RegisterTemplateResponse,
    TemplateVariableInput,
    TestGenerateRequest,
)
from services.bulk_document_generator import BulkDocumentGenerator
from services.docx_replacement_engine import DocxReplacementEngine
from services.input_cache import InputCacheService
from services.llm_marker_identifier import LLMMarkerIdentifier
from services.text_extractor import TextExtractor

router = APIRouter(tags=["template_registration"])

PROJECT_ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_UPLOAD_DIR = PROJECT_ROOT / "templates" / "registered"
TEMPLATE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

MARKER_PATTERN = re.compile(r"\{\{([^{}]+)\}\}")

text_extractor = TextExtractor()
identifier = LLMMarkerIdentifier()
docx_engine = DocxReplacementEngine()
input_cache_service = InputCacheService()


def _supported_extension(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in {"docx", "hwp", "hwpx"}:
        raise HTTPException(status_code=400, detail="Only .docx, .hwp, .hwpx templates are supported.")
    return ext


def _extract_existing_markers(text: str) -> list[str]:
    markers = sorted({matched.strip() for matched in MARKER_PATTERN.findall(text) if matched.strip()})
    return [f"{{{{{marker}}}}}" for marker in markers]


def _sanitize_filename(filename: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', "_", filename).strip()
    return cleaned or "template.docx"


def _insert_markers_into_docx(file_bytes: bytes, variable_rows: list[TemplateVariableInput]) -> bytes:
    document = Document(BytesIO(file_bytes))

    replacements: list[tuple[str, str]] = []
    for row in variable_rows:
        source_text = (row.text or "").strip()
        marker_name = (row.marker_name or "").strip()
        if not source_text or not marker_name:
            continue
        replacements.append((source_text, f"{{{{{marker_name}}}}}"))

    if not replacements:
        return file_bytes

    def replace_paragraph(paragraph) -> None:
        if not paragraph.runs:
            return
        full_text = "".join(run.text for run in paragraph.runs)
        next_text = full_text
        for source, marker in replacements:
            if source in next_text:
                next_text = next_text.replace(source, marker)
        if next_text == full_text:
            return
        # Preserve first run style and collapse text into first run.
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = next_text

    for paragraph in document.paragraphs:
        replace_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph(paragraph)

    for section in document.sections:
        for header in (section.header, section.first_page_header):
            if header:
                for paragraph in header.paragraphs:
                    replace_paragraph(paragraph)
        for footer in (section.footer, section.first_page_footer):
            if footer:
                for paragraph in footer.paragraphs:
                    replace_paragraph(paragraph)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


@router.post("/api/templates/analyze", response_model=AnalyzeTemplateResponse)
async def analyze_template(
    file: UploadFile = File(...),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Template file name is required.")

    ext = _supported_extension(file.filename)
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded template is empty.")

    try:
        extracted_text = text_extractor.extract(file_bytes, file.filename)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    existing_markers = set(_extract_existing_markers(extracted_text))
    if ext == "docx":
        for marker in docx_engine.extract_markers(file_bytes):
            existing_markers.add(f"{{{{{marker}}}}}")

    identified = await identifier.identify_markers(extracted_text)
    candidate_rows: list[MarkerCandidate] = []
    seen: set[tuple[str, str]] = set()
    for row in identified:
        marker = str(row.get("marker", "")).strip()
        if not marker:
            continue
        text = str(row.get("text", "")).strip()
        key = (text, marker)
        if key in seen:
            continue
        seen.add(key)
        candidate_rows.append(
            MarkerCandidate(
                text=text,
                marker=marker,
                source=str(row.get("source", "manual")).lower(),
                confidence=float(row.get("confidence", 0.5) or 0.5),
                existing=f"{{{{{marker}}}}}" in existing_markers,
            )
        )

    return AnalyzeTemplateResponse(
        extracted_text=extracted_text,
        identified_markers=candidate_rows,
        existing_markers=sorted(existing_markers),
    )


@router.post("/api/templates/register", response_model=RegisterTemplateResponse)
async def register_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    document_type: str = Form(...),
    variables: str = Form(...),
    db: Session = Depends(get_db),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Template file name is required.")
    ext = _supported_extension(file.filename)

    try:
        loaded_vars = json.loads(variables or "[]")
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail="`variables` must be valid JSON.") from exc
    if not isinstance(loaded_vars, list):
        raise HTTPException(status_code=400, detail="`variables` must be a JSON list.")

    variable_rows: list[TemplateVariableInput] = []
    for idx, item in enumerate(loaded_vars):
        if not isinstance(item, dict):
            continue
        payload = dict(item)
        if "display_order" not in payload:
            payload["display_order"] = idx
        try:
            variable_rows.append(TemplateVariableInput.model_validate(payload))
        except Exception:
            continue

    if not variable_rows:
        raise HTTPException(status_code=400, detail="At least one variable mapping is required.")

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded template is empty.")

    extracted_text = text_extractor.extract(file_bytes, file.filename)
    existing_markers = _extract_existing_markers(extracted_text)
    if ext == "docx" and not existing_markers:
        file_bytes = _insert_markers_into_docx(file_bytes, variable_rows)

    stamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    stored_name = f"{stamp}_{uuid4().hex}_{_sanitize_filename(file.filename)}"
    stored_path = TEMPLATE_UPLOAD_DIR / stored_name
    stored_path.write_bytes(file_bytes)

    relative_path = str(stored_path.relative_to(PROJECT_ROOT)).replace("\\", "/")
    marker_names = [row.marker_name for row in variable_rows]

    template = DocumentTemplate(
        name=name.strip(),
        category=document_type.strip(),
        file_path=relative_path,
        description="Registered via template registration wizard",
        variables=json.dumps(marker_names, ensure_ascii=False),
        custom_data="{}",
    )
    db.add(template)
    db.flush()

    for row in variable_rows:
        db.add(
            TemplateVariable(
                template_id=template.id,
                marker_name=row.marker_name,
                display_label=row.display_label,
                source_type=row.source_type,
                source_field=row.source_field,
                default_value=row.default_value,
                is_required=row.is_required,
                display_order=row.display_order,
            )
        )

    db.commit()
    db.refresh(template)

    return RegisterTemplateResponse(
        template_id=template.id,
        name=template.name,
        document_type=template.category,
        variable_count=len(variable_rows),
        file_path=template.file_path or "",
    )


@router.post("/api/templates/test-generate")
def test_generate_template(
    body: TestGenerateRequest,
    db: Session = Depends(get_db),
):
    template = db.get(DocumentTemplate, body.template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found.")

    cached = input_cache_service.get_last_inputs(db, body.fund_id, body.template_id)
    manual_vars = dict(cached)
    manual_vars.update({str(k): str(v) for k, v in body.manual_vars.items()})

    manual_marker_names = [
        row.marker_name
        for row in getattr(template, "template_variables", [])
        if (row.source_type or "").lower() == "manual"
    ]
    if manual_marker_names:
        filtered_manual = {
            marker: manual_vars.get(marker, "")
            for marker in manual_marker_names
            if marker in manual_vars
        }
    else:
        filtered_manual = dict(manual_vars)

    generator = BulkDocumentGenerator()
    try:
        preview_bytes = generator.preview_one(
            db=db,
            fund_id=body.fund_id,
            template_id=body.template_id,
            lp_id=body.lp_id,
            investment_id=body.investment_id,
            extra_vars=manual_vars,
        )
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    input_cache_service.save_inputs(
        db=db,
        fund_id=body.fund_id,
        template_id=body.template_id,
        manual_vars=filtered_manual,
    )

    return StreamingResponse(
        iter([preview_bytes]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": (
                f"attachment; filename*=UTF-8''preview_{template.name}_{datetime.utcnow():%Y%m%d}.docx"
            )
        },
    )


@router.get("/api/templates/input-cache", response_model=InputCacheResponse)
def get_template_input_cache(
    fund_id: int = Query(..., ge=1),
    template_id: int = Query(..., ge=1),
    db: Session = Depends(get_db),
):
    return InputCacheResponse(
        fund_id=fund_id,
        template_id=template_id,
        inputs=input_cache_service.get_last_inputs(db, fund_id, template_id),
    )
